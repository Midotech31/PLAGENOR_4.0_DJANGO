"""High-value document formatting and conversion regression tests."""

import sys
import tempfile
from pathlib import Path
from types import ModuleType, SimpleNamespace
from unittest.mock import MagicMock, patch

from django.test import SimpleTestCase, TestCase
from docx import Document
from docx.shared import Pt


class LegacyDocumentFormattingTests(SimpleTestCase):
    def test_header_footer_styles_and_semantic_anchors_are_idempotent(self):
        from documents.docx_helpers import (
            _find_anchor_for_position, _find_section_end, _is_heading_shaped,
            _is_section_heading, _matches_anchor, add_brand_footer,
            ensure_institutional_header, style_brand_table,
        )

        doc = Document()
        title = doc.add_heading('PLAGENOR', level=1)
        requester = doc.add_heading('1. Informations du demandeur', level=2)
        doc.add_paragraph('Requester details')
        samples = doc.add_heading('2. Tableau des échantillons', level=2)
        doc.add_paragraph('Samples')
        next_heading = doc.add_heading('3. Validation', level=2)
        footer_line = doc.add_paragraph('Document généré automatiquement')

        self.assertTrue(_is_heading_shaped(title))
        self.assertFalse(_is_section_heading(title))
        self.assertTrue(_matches_anchor(requester, ('demandeur',)))
        self.assertIs(_find_section_end(doc, samples)._element, next_heading._element)
        self.assertEqual(_find_anchor_for_position(doc, 'TOP')[1], 'after')
        self.assertIs(
            _find_anchor_for_position(doc, 'AFTER_REQUESTER')[0]._element,
            samples._element,
        )
        self.assertEqual(_find_anchor_for_position(doc, 'BOTTOM'), (None, 'end'))
        self.assertIs(
            _find_anchor_for_position(doc, 'BEFORE_FOOTER')[0]._element,
            footer_line._element,
        )
        self.assertEqual(_find_anchor_for_position(doc, 'UNKNOWN'), (None, 'end'))

        plain = doc.add_paragraph('Bold heading')
        run = plain.runs[0]
        run.bold = True
        run.font.size = Pt(14)
        self.assertTrue(_is_heading_shaped(plain))

        ensure_institutional_header(doc, Path('/definitely/missing/banner.png'))
        ensure_institutional_header(doc)
        header_xml = doc.sections[0].header._element.xml
        ensure_institutional_header(doc)
        self.assertEqual(doc.sections[0].header._element.xml, header_xml)

        add_brand_footer(doc)
        footer_xml = doc.sections[0].footer._element.xml
        add_brand_footer(doc)
        self.assertEqual(doc.sections[0].footer._element.xml, footer_xml)
        self.assertIn('NUMPAGES', footer_xml)

        empty_table = doc.add_table(rows=0, cols=2)
        style_brand_table(empty_table)

    def test_legacy_labels_samples_and_parameter_questions_are_filled(self):
        from documents.docx_helpers import (
            _build_param_label_index, _canonical_for, _fuzzy_pick,
            apply_legacy_label_substitution, populate_legacy_param_questions,
            populate_legacy_sample_table,
        )

        doc = Document()
        doc.add_paragraph('Nom et prénom : * Nom complet du demandeur')
        doc.add_paragraph('Laboratoire : * Cliquez ici')
        doc.add_paragraph('…….../2026/IBTIKAR/PLAGENOR/ESSBO')
        table = doc.add_table(rows=4, cols=4)
        for index, label in enumerate(
            ('N°', 'Code échantillon', "Source d'isolement", "Date d'isolement")
        ):
            table.cell(0, index).text = label
        table.cell(1, 0).paragraphs[0].add_run('old')
        nested = doc.add_table(rows=1, cols=1)
        nested.cell(0, 0).text = 'Adresse e-mail : * exemple'
        doc.sections[0].header.paragraphs[0].text = 'Téléphone : * exemple'
        doc.sections[0].footer.paragraphs[0].text = 'Titre du projet : * exemple'

        field_map = {
            'FULL_NAME': 'Marie Curie', 'LABORATORY': 'PLAGENOR',
            'EMAIL': 'marie@example.test', 'PHONE': '0550000000',
            'TITLE': 'Projet génomique', 'DISPLAY_ID': 'IBT-2026-0042',
        }
        apply_legacy_label_substitution(doc, field_map)
        all_xml = doc._element.xml + ''.join(
            section.header._element.xml + section.footer._element.xml
            for section in doc.sections
        )
        self.assertIn('Marie Curie', all_xml)
        self.assertIn('IBT-2026-0042', all_xml)
        self.assertIn('marie@example.test', all_xml)

        request_obj = SimpleNamespace(
            service=SimpleNamespace(code='SEQ-PHASE2'),
            sample_table=[
                {'sample_code': 'S-01', 'sample_origin': 'Sol',
                 'isolation_date': '2026-08-20'},
                {'sample_code': 'S-02', 'sample_origin': 'Eau',
                 'isolation_date': '2026-08-21'},
                None,
            ],
            service_params={'kit_type': 'Premium', 'replicates': 3,
                            'paired_end': True, 'empty': ''},
        )
        registry = {
            'sample_table': {'columns': [
                {'name': 'sample_code', 'label': 'Code échantillon'},
                {'name': 'sample_origin', 'label': "Source d'isolement"},
                {'name': 'isolation_date', 'label': "Date d'isolement"},
            ]},
            'parameters': [
                {'name': 'kit_type', 'label_fr': 'Type de kit PCR',
                 'label_en': 'PCR kit type'},
                {'name': 'replicates', 'label': 'Nombre de répétitions'},
                {'name': 'paired_end', 'label': 'Lectures appariées'},
            ],
        }
        with patch('core.registry.get_service_def', return_value=registry):
            populate_legacy_sample_table(doc, request_obj)
            question = doc.add_paragraph(
                'Type de kit PCR souhaité : Choisissez un élément')
            personal = doc.add_paragraph(
                'Nom et prénom du demandeur : Cliquez ou appuyez ici')
            unknown = doc.add_paragraph(
                'Couleur préférée : Choisissez un élément')
            populate_legacy_param_questions(doc, request_obj)
            index = _build_param_label_index(
                request_obj.service_params, request_obj.service.code)

        flattened = '\n'.join(
            cell.text for row in table.rows for cell in row.cells)
        self.assertIn('S-01', flattened)
        self.assertIn('S-02', flattened)
        self.assertIn('2026-08-21', flattened)
        self.assertEqual(len(table.rows), 3)
        self.assertIn('Premium', question.text)
        self.assertIn('Cliquez', personal.text)
        self.assertIn('Choisissez', unknown.text)
        self.assertTrue(index)
        self.assertIsNone(_canonical_for(''))
        self.assertIsNone(_fuzzy_pick('', index, set()))


class LibreOfficeBackendTests(SimpleTestCase):
    def setUp(self):
        self.temp_dir = tempfile.TemporaryDirectory(prefix='plagenor-phase2-pdf-')
        self.root = Path(self.temp_dir.name)
        self.docx = self.root / 'input.docx'
        self.docx.write_bytes(b'docx')

    def tearDown(self):
        self.temp_dir.cleanup()

    def test_spawn_handles_unexpected_crash_and_binary_alias(self):
        from documents.pdf_converter import _convert_via_spawn, _soffice_bin

        with patch('documents.pdf_converter.shutil.which', side_effect=[None, '/bin/libreoffice']):
            self.assertEqual(_soffice_bin(), '/bin/libreoffice')
        with patch('documents.pdf_converter._soffice_bin', return_value='/bin/soffice'), \
             patch('documents.pdf_converter.subprocess.run', side_effect=OSError('crash')):
            self.assertEqual(_convert_via_spawn(self.docx, self.root), self.docx)

    def test_daemon_lifecycle_start_and_shutdown(self):
        from documents.pdf_converter import _LibreOfficeDaemon, _UnoConversionError

        daemon = _LibreOfficeDaemon()
        self.assertFalse(daemon._alive())
        with patch('documents.pdf_converter._soffice_bin', return_value=None):
            with self.assertRaises(_UnoConversionError):
                daemon._start()

        process = MagicMock()
        process.poll.return_value = None
        with patch('documents.pdf_converter._soffice_bin', return_value='/bin/soffice'), \
             patch('documents.pdf_converter.subprocess.Popen', return_value=process) as popen:
            daemon._start()
        self.assertTrue(daemon._alive())
        self.assertIn('--headless', popen.call_args.args[0])
        daemon.shutdown()
        process.terminate.assert_called_once()
        self.assertIsNone(daemon._proc)

        failing_process = MagicMock()
        failing_process.wait.side_effect = RuntimeError('stuck')
        daemon._proc = failing_process
        daemon.shutdown()
        failing_process.kill.assert_called_once()

    def test_uno_conversion_success_and_fallback_paths(self):
        from documents.pdf_converter import (
            _LibreOfficeDaemon, _UnoConversionError, _convert_via_uno,
        )

        class PropertyValue:
            pass

        uno = ModuleType('uno')
        uno.systemPathToFileUrl = lambda value: f'file://{value}'
        com = ModuleType('com')
        star = ModuleType('com.sun.star')
        beans = ModuleType('com.sun.star.beans')
        beans.PropertyValue = PropertyValue
        modules = {
            'uno': uno, 'com': com, 'com.sun': ModuleType('com.sun'),
            'com.sun.star': star, 'com.sun.star.beans': beans,
        }

        output = self.root / 'input.pdf'
        fake_document = MagicMock()
        fake_document.storeToURL.side_effect = lambda *_args: output.write_bytes(b'%PDF')
        desktop = MagicMock()
        desktop.loadComponentFromURL.return_value = fake_document
        daemon = _LibreOfficeDaemon()
        with patch.dict(sys.modules, modules), \
             patch.object(daemon, '_ensure_running'), \
             patch.object(daemon, '_connect_desktop', return_value=desktop):
            self.assertEqual(daemon.convert(self.docx, self.root), output)
        fake_document.close.assert_called_once_with(False)

        output.unlink()
        desktop.loadComponentFromURL.return_value = None
        with patch.dict(sys.modules, modules), \
             patch.object(daemon, '_ensure_running'), \
             patch.object(daemon, '_connect_desktop', return_value=desktop):
            with self.assertRaises(_UnoConversionError):
                daemon.convert(self.docx, self.root)

        with patch.dict(sys.modules, {'uno': None}), \
             patch('documents.pdf_converter._convert_via_spawn', return_value=self.docx) as spawn:
            self.assertEqual(_convert_via_uno(self.docx, self.root), self.docx)
            spawn.assert_called_once()

        with patch.dict(sys.modules, {'uno': uno}), \
             patch('documents.pdf_converter._get_daemon') as get_daemon, \
             patch('documents.pdf_converter._convert_via_spawn', return_value=self.docx) as spawn:
            get_daemon.return_value.convert.side_effect = _UnoConversionError('expected')
            self.assertEqual(_convert_via_uno(self.docx, self.root), self.docx)
            get_daemon.return_value.convert.side_effect = RuntimeError('unexpected')
            self.assertEqual(_convert_via_uno(self.docx, self.root), self.docx)
            self.assertEqual(spawn.call_count, 2)


class ProgrammaticDocumentBuilderTests(TestCase):
    @classmethod
    def setUpTestData(cls):
        from accounts.models import User
        from core.models import Request, Service

        cls.requester = User.objects.create_user(
            username='phase2-doc-requester', password='test-password',
            role='REQUESTER', first_name='Marie', last_name='Curie',
            email='marie@example.test', phone='0550000000',
            organization='ESSBO', laboratory='PLAGENOR',
            supervisor='Pr Test', student_level='Doctorat',
            preferred_language='en',
        )
        analyst_user = User.objects.create_user(
            username='phase2-doc-analyst', password='test-password',
            role='MEMBER', first_name='Ada', last_name='Lovelace',
            email='ada@example.test',
        )
        cls.service = Service.objects.create(
            code='PHASE2-DOC', name='Séquençage',
            description='Analyse approfondie', turnaround_days=7,
            channel_availability='BOTH', ibtikar_price=1000,
        )
        cls.request_obj = Request.objects.create(
            display_id='PHASE2-DOC-001', channel='IBTIKAR',
            status='VALIDATION_FINANCE', service=cls.service,
            requester=cls.requester, assigned_to=analyst_user.member_profile,
            title='Projet génomique', description='Description scientifique',
            urgency='Urgent', budget_amount=5000,
            declared_ibtikar_balance=10000, admin_validated_price=4500,
            service_params={'analysis_mode': 'duplicate', 'pathogenic': True,
                            'analysis_frame': 'Doctorat'},
            sample_table=[
                {'sample_code': 'S1', 'organism_type': 'Bactérie'},
                {'sample_code': 'S2', 'organism_type': 'Levure'},
            ],
        )

    def test_programmatic_fallback_builders_render_complete_documents(self):
        from documents.generators import (
            _build_platform_note_programmatic, _build_quote_programmatic,
            _build_reception_form_programmatic, _detableize, build_field_map,
        )

        field_map = build_field_map(self.request_obj)
        field_map['QUOTE_NUMBER'] = 'GENOCLAB-DEV-2026-0001'
        registry = {
            'pricing': {
                'base_price': {'pathogenic': 1000, 'non_pathogenic': 800},
                'multipliers': {'duplicate': 2},
            },
        }
        with patch('core.pricing.resolve_cost', return_value={
            'total': 4000,
            'breakdown': {
                'base_price': 1000, 'multiplier': 2,
                'multiplier_key': 'duplicate', 'pathogenic': True,
            },
        }), patch('core.registry.get_service_def', return_value=registry):
            platform = _build_platform_note_programmatic(
                self.request_obj, field_map)
        quote = _build_quote_programmatic(self.request_obj, field_map)
        reception = _build_reception_form_programmatic(
            self.request_obj, field_map)

        platform_text = '\n'.join(p.text for p in platform.paragraphs)
        self.assertIn('Marie Curie', platform_text)
        self.assertIn('Réajusté par administration', platform_text)
        self.assertIn('2 échantillons', platform_text)
        self.assertGreater(len(quote.tables), 0)
        self.assertGreater(len(reception.tables), 0)

        _detableize(quote)
        self.assertEqual(len(quote.tables), 0)
        quote_text = '\n'.join(p.text for p in quote.paragraphs)
        self.assertIn('Total TTC', quote_text)

        three_col = Document()
        table = three_col.add_table(rows=2, cols=3)
        table.cell(0, 0).text = 'A'
        table.cell(0, 1).text = 'B'
        table.cell(0, 2).text = 'C'
        _detableize(three_col)
        self.assertIn('A — B — C', '\n'.join(p.text for p in three_col.paragraphs))

    def test_tariff_fallbacks_and_document_block_positions(self):
        from documents.generators import (
            _inject_document_blocks, _render_tariff_breakdown,
        )
        from documents.models import DocumentBlock

        # A list-shaped pricing breakdown exercises the database-tier fallback
        # to registry pricing and request parameters.
        tariff_doc = Document()
        registry = {
            'pricing': {
                'base_price': {'pathogenic': 1100, 'default': 900},
                'multipliers': {'duplicate': 2},
            },
        }
        with patch('core.pricing.resolve_cost', return_value={
            'total': 4400, 'breakdown': [{'label': 'tier'}],
        }), patch('core.registry.get_service_def', return_value=registry):
            _render_tariff_breakdown(tariff_doc, self.request_obj)
        tariff_text = '\n'.join(p.text for p in tariff_doc.paragraphs)
        self.assertIn('1 100 DA', tariff_text)
        self.assertIn('× 2', tariff_text)

        # Pricing/registry failures must still yield a readable document.
        no_price = SimpleNamespace(
            service=self.service, service_params={}, sample_table=[],
            channel='IBTIKAR', urgency='Normal', budget_amount='indisponible',
            admin_validated_price=None,
        )
        fallback_doc = Document()
        with patch('core.pricing.resolve_cost', side_effect=RuntimeError('pricing')), \
             patch('core.registry.get_service_def', side_effect=RuntimeError('registry')):
            _render_tariff_breakdown(fallback_doc, no_price)
        self.assertIn('indisponible', '\n'.join(p.text for p in fallback_doc.paragraphs))

        for priority, position in enumerate(
            ('TOP', 'AFTER_REQUESTER', 'AFTER_SAMPLES',
             'BEFORE_FOOTER', 'BOTTOM')
        ):
            DocumentBlock.objects.create(
                template_type='PLATFORM_NOTE', position=position,
                language='fr', priority=priority,
                title=f'{position} — {{{{DISPLAY_ID}}}}',
                body=f'Contenu {position}\n\nDeuxième paragraphe',
            )
        DocumentBlock.objects.create(
            template_type='PLATFORM_NOTE', position='BOTTOM',
            language='fr', title='', body='', priority=99,
        )
        doc = Document()
        doc.add_heading('Titre', level=1)
        doc.add_heading('Demandeur', level=2)
        doc.add_paragraph('Informations')
        doc.add_heading('Échantillons', level=2)
        doc.add_paragraph('Données')
        doc.add_heading('Validation', level=2)
        doc.add_paragraph('Document généré automatiquement')
        _inject_document_blocks(doc, 'PLATFORM_NOTE', self.request_obj)
        rendered = '\n'.join(p.text for p in doc.paragraphs)
        for position in ('TOP', 'AFTER_REQUESTER', 'AFTER_SAMPLES',
                         'BEFORE_FOOTER', 'BOTTOM'):
            self.assertIn(f'{position} — PHASE2-DOC-001', rendered)
