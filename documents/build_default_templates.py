"""One-shot builder for the four generic DOCX templates.

The originally-shipped ``quote_template.docx`` and
``reception_form_template.docx`` used ``[Service]``-style markers that no
generator ever substituted, leaking raw template syntax into customer-
facing documents. ``platform_note_template.docx`` was worse — a hardcoded
EGTP-Seq02 narrative with no placeholders at all, so every generated
Platform Note showed the same Seq02 text regardless of request.

This module rebuilds the three problem templates with consistent
``{{KEY}}`` placeholders, A4 page size, 1" margins, the institutional
banner in the header, and a discreet generation footer. The IBTIKAR
template is left alone — it already uses ``{{KEY}}`` correctly.

Run via ``python manage.py shell -c "from documents.build_default_templates import build_all; build_all()"``.
The output overwrites ``documents/docx_templates/*.docx``; the originals
are backed up with a ``.bak.docx`` suffix the first time.
"""
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentType
from docx.shared import Pt

from documents.docx_helpers import apply_house_style, ensure_institutional_header


TEMPLATE_DIR = Path(__file__).resolve().parent / 'docx_templates'


def _backup(path: Path) -> None:
    if path.exists():
        backup = path.with_suffix('.bak.docx')
        if not backup.exists():
            shutil.copy2(str(path), str(backup))


def _kv_table(doc: DocumentType, rows) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = 'Light Grid Accent 1'
    for i, (label, value) in enumerate(rows):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value


def build_platform_note_template() -> Path:
    doc = Document()
    apply_house_style(doc)

    doc.add_heading('NOTE DE PLATEFORME — PLAGENOR', level=1)
    doc.add_paragraph("ESSBO — École Supérieure en Sciences Biologiques d'Oran")
    doc.add_paragraph('Référence : {{DISPLAY_ID}}')
    doc.add_paragraph("Date d'émission : {{DATETIME}}")
    doc.add_paragraph('')

    doc.add_heading('Demandeur', level=2)
    _kv_table(doc, [
        ('Nom complet', '{{FULL_NAME}}'),
        ('Établissement', '{{ETABLISSEMENT}}'),
        ('Laboratoire', '{{LABORATORY}}'),
        ('Niveau / fonction', '{{STUDENT_LEVEL}}'),
        ('Directeur de recherche', '{{SUPERVISOR}}'),
        ('Email', '{{EMAIL}}'),
        ('Téléphone', '{{PHONE}}'),
    ])

    doc.add_heading('Service demandé', level=2)
    _kv_table(doc, [
        ('Code', '{{SERVICE_CODE}}'),
        ('Intitulé', '{{SERVICE_NAME}}'),
        ('Description', '{{SERVICE_DESCRIPTION}}'),
        ('Délai (jours ouvrables)', '{{SERVICE_TURNAROUND}}'),
        ('Canal', '{{CHANNEL}}'),
        ('Urgence', '{{URGENCY}}'),
    ])

    doc.add_heading('Détails de la demande', level=2)
    doc.add_paragraph('Titre : {{TITLE}}')
    doc.add_paragraph('Description : {{DESCRIPTION}}')
    doc.add_paragraph('Paramètres : {{SERVICE_PARAMS}}')
    doc.add_paragraph('Échantillons : {{SAMPLE_TABLE}}')

    doc.add_heading('Décompte budgétaire IBTIKAR', level=2)
    doc.add_paragraph('Budget annuel par étudiant : 200 000 DZD')
    doc.add_paragraph('Montant de cette prestation : {{BUDGET_AMOUNT}}')
    doc.add_paragraph('Solde IBTIKAR déclaré : {{IBTIKAR_BALANCE}}')

    doc.add_heading('Assignation', level=2)
    doc.add_paragraph('Analyste : {{ASSIGNED_ANALYST}}')
    doc.add_paragraph('Email analyste : {{ANALYST_EMAIL}}')
    doc.add_paragraph('Rendez-vous : {{APPOINTMENT_DATE}}')

    _add_footer(doc)
    ensure_institutional_header(doc)

    path = TEMPLATE_DIR / 'platform_note_template.docx'
    _backup(path)
    doc.save(str(path))
    return path


def build_quote_template() -> Path:
    doc = Document()
    apply_house_style(doc)

    doc.add_heading('DEVIS — GENOCLAB', level=1)
    doc.add_paragraph("ESSBO — École Supérieure en Sciences Biologiques d'Oran")
    doc.add_paragraph('N° Devis : {{QUOTE_NUMBER}}')
    doc.add_paragraph('Référence demande : {{DISPLAY_ID}}')
    doc.add_paragraph('Date : {{DATE}}')
    doc.add_paragraph('')

    doc.add_heading('Client', level=2)
    _kv_table(doc, [
        ('Nom', '{{CLIENT_NAME}}'),
        ('Organisation', '{{ORGANIZATION}}'),
        ('Laboratoire', '{{LABORATORY}}'),
        ('Email', '{{CLIENT_EMAIL}}'),
        ('Téléphone', '{{PHONE}}'),
    ])

    doc.add_heading('Prestations', level=2)
    table = doc.add_table(rows=2, cols=4)
    table.style = 'Light Grid Accent 1'
    for j, h in enumerate(['Description', 'Quantité', 'Prix unitaire', 'Total']):
        table.rows[0].cells[j].text = h
    table.rows[1].cells[0].text = '{{SERVICE_NAME}}'
    table.rows[1].cells[1].text = '1'
    table.rows[1].cells[2].text = '{{SUBTOTAL_HT}}'
    table.rows[1].cells[3].text = '{{SUBTOTAL_HT}}'

    doc.add_paragraph('')
    summary = doc.add_table(rows=3, cols=2)
    summary.style = 'Light Grid Accent 1'
    summary.rows[0].cells[0].text = 'Sous-total HT'
    summary.rows[0].cells[1].text = '{{SUBTOTAL_HT}}'
    summary.rows[1].cells[0].text = 'TVA ({{VAT_RATE}})'
    summary.rows[1].cells[1].text = '{{VAT_AMOUNT}}'
    summary.rows[2].cells[0].text = 'Total TTC'
    summary.rows[2].cells[1].text = '{{TOTAL_TTC}}'

    _add_footer(doc)
    ensure_institutional_header(doc)

    path = TEMPLATE_DIR / 'quote_template.docx'
    _backup(path)
    doc.save(str(path))
    return path


def build_reception_form_template() -> Path:
    doc = Document()
    apply_house_style(doc)

    doc.add_heading("FICHE DE RÉCEPTION D'ÉCHANTILLONS", level=1)
    doc.add_paragraph('PLAGENOR — ESSBO')
    doc.add_paragraph('Référence : {{DISPLAY_ID}}')
    doc.add_paragraph('Code de suivi : {{TRACKING_CODE}}')
    doc.add_paragraph('')

    _kv_table(doc, [
        ('Service', '{{SERVICE_NAME}}'),
        ('Canal', '{{CHANNEL}}'),
        ('Urgence', '{{URGENCY}}'),
        ('Date de RDV', '{{APPOINTMENT_DATE}}'),
        ('Analyste assigné', '{{ASSIGNED_ANALYST}}'),
        ('Date de soumission', '{{SUBMISSION_DATE}}'),
    ])

    doc.add_heading('Déposant', level=2)
    _kv_table(doc, [
        ('Nom', '{{FULL_NAME}}'),
        ('Email', '{{EMAIL}}'),
        ('Téléphone', '{{PHONE}}'),
        ('Établissement', '{{ETABLISSEMENT}}'),
        ('Laboratoire', '{{LABORATORY}}'),
    ])

    doc.add_heading('Échantillons soumis', level=2)
    doc.add_paragraph('{{SAMPLE_TABLE}}')

    doc.add_heading('Contrôle à la réception', level=2)
    _kv_table(doc, [
        ('Date de réception', '___ / ___ / ______'),
        ("Nombre d'échantillons reçus", '____________'),
        ('État des échantillons', '☐ Bon   ☐ Acceptable   ☐ Dégradé'),
        ('Observations', ''),
    ])

    doc.add_paragraph('')
    doc.add_paragraph('Signature du réceptionniste : ________________________')
    doc.add_paragraph('Signature du déposant : ________________________')

    _add_footer(doc)
    ensure_institutional_header(doc)

    path = TEMPLATE_DIR / 'reception_form_template.docx'
    _backup(path)
    doc.save(str(path))
    return path


def _add_footer(doc: DocumentType) -> None:
    section = doc.sections[0]
    footer = section.footer
    if footer.paragraphs:
        paragraph = footer.paragraphs[0]
    else:
        paragraph = footer.add_paragraph()
    paragraph.text = ''
    run = paragraph.add_run('Document généré automatiquement par PLAGENOR 4.0 · ESSBO')
    run.font.size = Pt(9)


def build_all():
    """Build all three programmatic templates. Returns the list of paths."""
    return [
        build_platform_note_template(),
        build_quote_template(),
        build_reception_form_template(),
    ]


if __name__ == '__main__':
    import os
    os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'plagenor.settings')
    import django
    django.setup()
    for p in build_all():
        print(f'wrote {p}')
