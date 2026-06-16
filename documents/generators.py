"""PLAGENOR document generators.

Each ``generate_*`` function returns a filesystem path to a finalised DOCX.
The path is then optionally rendered to PDF by ``documents.views`` via the
``pdf_converter`` module.

Phase 3.7 refactor highlights:
* Single :func:`build_field_map` is the source of truth for every placeholder
  the system supports — all four generators draw from the same dict, so an
  IBTIKAR form and a Platform Note never disagree on what ``{{FULL_NAME}}``
  resolves to.
* Placeholder substitution goes through :func:`docx_helpers.replace_placeholders`,
  which is run-preserving (the old ``paragraph.text = ...`` collapsed every
  run into one) and which only touches ``{{KEY}}``-wrapped tokens. The
  previous bare-key replacement caused the
  ``{{APPOINTMENT_24/05/2026}}`` corruption seen in reception forms.
* Programmatic-fallback generation applies the house style (A4, 1" margins,
  11 pt body) and injects the institutional logo banner into the header so
  that fallback documents look as branded as the IBTIKAR templates.
* Admin-editable :class:`documents.models.DocumentBlock` rows are injected
  at semantic positions (``TOP``, ``AFTER_REQUESTER``, ``AFTER_SAMPLES``,
  ``BEFORE_FOOTER``, ``BOTTOM``) so SuperAdmins can add notices/disclaimers
  globally or per-service without re-uploading the DOCX.
"""
from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path
from typing import Optional

from django.conf import settings
from docx import Document
from docx.document import Document as DocumentType
from docx.shared import Pt

from documents.docx_helpers import (
    _find_anchor_for_position,
    add_brand_footer,
    add_paragraph_after,
    add_paragraph_before,
    apply_house_style,
    apply_legacy_label_substitution,
    ensure_institutional_header,
    populate_legacy_param_questions,
    populate_legacy_sample_table,
    replace_placeholders,
    strip_unresolved_placeholders,
    style_brand_table,
)

_PLACEHOLDER_BARE_RE = re.compile(r'\{\{[A-Z0-9_]+\}\}')


IBTIKAR_TEMPLATE_MAP = {
    'EGTP-CAN': 'egtp_can.docx',
    'EGTP-IMT': 'egtp_imt.docx',
    'EGTP-PCR': 'egtp_pcr.docx',
    'EGTP-Lyoph': 'egtp_lyoph.docx',
    'EGTP-PS': 'egtp_ps.docx',
    'EGTP-Seq02': 'egtp_seq02.docx',
    'EGTP-SeqS': 'egtp_seqs.docx',
    'EGTP-GDE': 'egtp_gde.docx',
    'EGTP-Illumina-Microbial-WGS': 'egtp_illumina_wgs.docx',
}


def _money(value, currency='DZD', placeholder='N/A'):
    if value is None:
        return placeholder
    try:
        return f"{value:,.0f} {currency}"
    except (TypeError, ValueError):
        return placeholder


def _money_2dp(value, currency='DZD', placeholder='N/A'):
    if value is None:
        return placeholder
    try:
        return f"{value:,.2f} {currency}"
    except (TypeError, ValueError):
        return placeholder


def _safe_attr(obj, attr, default=''):
    try:
        return getattr(obj, attr, default) or default
    except Exception:
        return default


def _assigned_name(request_obj, placeholder='Non assigné'):
    try:
        if request_obj.assigned_to and request_obj.assigned_to.user:
            user = request_obj.assigned_to.user
            return user.get_full_name() or user.username or placeholder
    except Exception:
        pass
    return placeholder


def _assigned_email(request_obj, placeholder=''):
    try:
        if request_obj.assigned_to and request_obj.assigned_to.user:
            return _safe_attr(request_obj.assigned_to.user, 'email', placeholder)
    except Exception:
        pass
    return placeholder


def _format_date(value, placeholder='Non défini'):
    if not value:
        return placeholder
    try:
        return value.strftime('%d/%m/%Y')
    except Exception:
        return placeholder


def _format_sample_table_text(sample_table):
    """Render the sample_table JSON as a readable multi-line string."""
    if not sample_table or not isinstance(sample_table, list):
        return ''
    lines = []
    for i, sample in enumerate(sample_table, 1):
        if not isinstance(sample, dict):
            continue
        non_empty = [(k, v) for k, v in sample.items() if v not in (None, '', [])]
        if not non_empty:
            continue
        lines.append(f"#{i} " + " | ".join(f"{k}: {v}" for k, v in non_empty))
    return '\n'.join(lines)


def _format_service_params(service_params):
    """Render the service_params JSON as a readable key: value block."""
    if not service_params or not isinstance(service_params, dict):
        return ''
    out = []
    for key, value in service_params.items():
        if value in (None, '', [], {}):
            continue
        clean = key.replace('param_', '').replace('_', ' ').strip().capitalize()
        out.append(f"{clean}: {value}")
    return '\n'.join(out)


def build_field_map(request_obj) -> dict[str, str]:
    """Single source of truth for placeholder values.

    Every key listed here is documented in the SuperAdmin "Modèles de
    documents" help screen; the four generators draw from the same dict
    so callers never see drift between IBTIKAR and Platform Note for the
    same request.

    Naming conventions:
    * REQUESTER_*  — fields off ``request.requester`` (User row). For a
      guest submission these fall back to ``request.guest_*``.
    * SERVICE_*    — fields off ``request.service`` (Service row).
    * REQUEST_*    — fields directly on the Request row.
    * QUOTE_*      — derived financial figures for GENOCLAB.
    * ASSIGNED_*   — about the analyst the request is assigned to.
    * DATE / DATETIME / *_DATE — formatted dates.

    Legacy short-names (FULL_NAME, EMAIL, PHONE, etc.) are also emitted so
    pre-existing templates keep working — both forms are valid.
    """
    req = request_obj
    requester = req.requester

    if requester is not None:
        full_name = requester.get_full_name() or requester.username
        email = _safe_attr(requester, 'email')
        phone = _safe_attr(requester, 'phone')
        organization = _safe_attr(requester, 'organization')
        laboratory = _safe_attr(requester, 'laboratory')
        supervisor = _safe_attr(requester, 'supervisor')
        student_level = _safe_attr(requester, 'student_level')
        username = _safe_attr(requester, 'username')
        ibtikar_id = _safe_attr(requester, 'ibtikar_id')
    else:
        full_name = req.guest_name or 'N/A'
        email = req.guest_email or ''
        phone = req.guest_phone or ''
        organization = ''
        laboratory = ''
        supervisor = ''
        student_level = ''
        username = ''
        ibtikar_id = ''

    service = req.service
    if service is not None:
        svc_code = service.code
        svc_name = service.name
        svc_description = service.description or ''
        svc_turnaround = str(service.turnaround_days) if service.turnaround_days else 'N/A'
        svc_type = _safe_attr(service, 'service_type', 'Analysis')
    else:
        svc_code = 'N/A'
        svc_name = 'N/A'
        svc_description = ''
        svc_turnaround = 'N/A'
        svc_type = ''

    quote_amount = float(req.quote_amount or 0)
    vat_rate = float(getattr(settings, 'VAT_RATE', 0.19) or 0)
    vat_amount = round(quote_amount * vat_rate, 2)
    total_ttc = round(quote_amount + vat_amount, 2)

    field_map = {
        # ----- Dates --------------------------------------------------------
        'DATE': datetime.now().strftime('%d/%m/%Y'),
        'DATETIME': datetime.now().strftime('%d/%m/%Y à %H:%M'),
        'SUBMISSION_DATE': _format_date(req.created_at, placeholder=datetime.now().strftime('%d/%m/%Y')),
        'APPOINTMENT_DATE': _format_date(req.appointment_date),
        'CURRENT_YEAR': str(datetime.now().year),

        # ----- Request IDs -------------------------------------------------
        'DISPLAY_ID': req.display_id or '',
        'REQUEST_ID': str(req.pk),
        'TRACKING_CODE': str(req.guest_token or req.display_id or req.pk),
        'IBTIKAR_EXTERNAL_CODE': _safe_attr(req, 'ibtikar_external_code'),

        # ----- Requester / Client ------------------------------------------
        'FULL_NAME': full_name,
        'REQUESTER_NAME': full_name,
        'CLIENT_NAME': full_name,
        'USERNAME': username,
        'EMAIL': email,
        'REQUESTER_EMAIL': email,
        'CLIENT_EMAIL': email,
        'PHONE': phone,
        'REQUESTER_PHONE': phone,
        'ETABLISSEMENT': organization,
        'ORGANIZATION': organization,
        'LABORATORY': laboratory,
        'SUPERVISOR': supervisor,
        'STUDENT_LEVEL': student_level,
        'GUEST_NAME': req.guest_name or '',
        'GUEST_EMAIL': req.guest_email or '',
        'GUEST_PHONE': req.guest_phone or '',
        'IBTIKAR_ID': ibtikar_id,            # User.ibtikar_id (registration-time DGRSDT ID)
        'REQUESTER_IBTIKAR_ID': ibtikar_id,  # alias

        # ----- Service ------------------------------------------------------
        'SERVICE_CODE': svc_code,
        'SERVICE_NAME': svc_name,
        'SERVICE_DESCRIPTION': svc_description,
        'SERVICE_TURNAROUND': svc_turnaround,
        'TURNAROUND': svc_turnaround,
        'SERVICE_TYPE': svc_type,

        # ----- Request details ---------------------------------------------
        'TITLE': req.title or '',
        'PROJECT_TITLE': req.title or '',
        'DESCRIPTION': req.description or '',
        'CHANNEL': req.channel or '',
        'URGENCY': req.urgency or '',
        'STATUS': req.status or '',
        # ``analysis_frame`` is a per-service YAML param the requester picks
        # online (e.g. "PFE Classique" / "1275" / "008" / "Projet de doctorat"
        # / "Autre"). Surfacing it in the field_map lets the legacy IBTIKAR
        # forms substitute their "Cadre de l'analyse" line.
        'ANALYSIS_FRAME': (req.service_params or {}).get('analysis_frame', ''),

        # ----- Financial ----------------------------------------------------
        'BUDGET_AMOUNT': _money(req.budget_amount),
        'IBTIKAR_BUDGET': _money(req.budget_amount),
        'IBTIKAR_BALANCE': _money(req.declared_ibtikar_balance),
        'FINAL_COST': _money(req.admin_validated_price, placeholder='En attente'),
        'QUOTE_AMOUNT': _money_2dp(quote_amount),
        'SUBTOTAL_HT': _money_2dp(quote_amount),
        'VAT_RATE': f"{vat_rate * 100:.0f}%",
        'VAT_AMOUNT': _money_2dp(vat_amount),
        'TVA': _money_2dp(vat_amount),  # legacy alias
        'TOTAL_TTC': _money_2dp(total_ttc),

        # ----- Assignment ---------------------------------------------------
        'ASSIGNED_ANALYST': _assigned_name(req),
        'ANALYST_NAME': _assigned_name(req),
        'ANALYST_EMAIL': _assigned_email(req),

        # ----- Free-text dumps ---------------------------------------------
        'SAMPLE_TABLE': _format_sample_table_text(req.sample_table),
        'SERVICE_PARAMS': _format_service_params(req.service_params),
    }
    return field_map


# DocumentBlock injection ---------------------------------------------------

def _resolve_block_language(request_obj) -> str:
    """Pick the language for DocumentBlock matching.

    Order of preference: requester's preferred_language → request channel
    default → settings.LANGUAGE_CODE. Validated against the supported
    language set so an unknown code can never reach the DB filter.
    """
    supported = {c for c, _ in getattr(settings, 'LANGUAGES', [])} or {'fr'}
    requester = getattr(request_obj, 'requester', None)
    if requester is not None:
        pref = getattr(requester, 'preferred_language', '') or ''
        if pref in supported:
            return pref
    default = getattr(settings, 'LANGUAGE_CODE', 'fr')
    return default if default in supported else 'fr'


def _resolve_block_text(text: str, field_map: dict) -> str:
    """Resolve ``{{KEY}}`` placeholders inside admin-authored block text.

    Mirrors the run-aware DOCX substitution behaviour for plain strings:
    only ``{{KEY}}``-wrapped tokens are touched (never bare keys), and
    unresolved tokens are stripped at the end so a stray placeholder
    never reaches the reader.
    """
    if not text:
        return ''
    for key, value in field_map.items():
        token = f'{{{{{key}}}}}'
        if token in text:
            text = text.replace(token, '' if value is None else str(value))
    return _PLACEHOLDER_BARE_RE.sub('', text)


def _inject_document_blocks(doc: DocumentType, template_type: str, request_obj) -> None:
    """Insert admin-editable DocumentBlocks at the configured positions.

    For each position label (TOP / AFTER_REQUESTER / AFTER_SAMPLES /
    BEFORE_FOOTER / BOTTOM) the helper resolves an anchor paragraph in
    the current document and inserts the matching blocks there in
    ``priority`` order. When no semantic anchor is found (e.g. a Quote
    has no "Échantillons" heading), the position falls back to BOTTOM
    so the block still appears.

    Placeholders inside the block's title and body
    (``{{FULL_NAME}}``, ``{{DISPLAY_ID}}``, ``{{DATE}}``, …) are
    resolved against the same field_map the template substitution uses,
    so admins author dynamic notices without escaping anything.

    Insertion happens after every other DOCX mutation so anchor lookup
    sees the final paragraph layout (including any banner-injection
    side-effects in headers — which don't appear in the body and so
    don't affect anchor matching).
    """
    from documents.models import DocumentBlock

    language = _resolve_block_language(request_obj)
    blocks = list(DocumentBlock.applicable_blocks(
        template_type=template_type,
        service=request_obj.service,
        language=language,
    ))
    if not blocks and language != settings.LANGUAGE_CODE:
        # Fall back to the default-language blocks if the active language
        # has nothing to say — same policy as the cms template tag.
        blocks = list(DocumentBlock.applicable_blocks(
            template_type=template_type,
            service=request_obj.service,
            language=settings.LANGUAGE_CODE,
        ))
    if not blocks:
        return

    field_map = build_field_map(request_obj)

    # Group blocks by position so we can resolve anchors once per group
    # and insert in deterministic order. Within a position group, blocks
    # are inserted in ``priority`` (then pk) order — the queryset is
    # already sorted that way.
    by_position: dict[str, list] = {}
    for block in blocks:
        by_position.setdefault(block.position, []).append(block)

    # Process positions in a stable presentation order so multiple
    # positions referenced in the same document land predictably.
    ordered_positions = ['TOP', 'AFTER_REQUESTER', 'AFTER_SAMPLES', 'BEFORE_FOOTER', 'BOTTOM']

    for position in ordered_positions:
        position_blocks = by_position.get(position, [])
        if not position_blocks:
            continue
        anchor, where = _find_anchor_for_position(doc, position)
        if anchor is None:
            # No semantic anchor → append at end.
            for block in position_blocks:
                _append_block_at_end(doc, block, field_map)
        else:
            # Insert each block above/below the anchor. For 'after' we keep
            # advancing the anchor so blocks land in the right order rather
            # than being reversed by sequential addnext calls.
            current_anchor = anchor
            for block in position_blocks:
                current_anchor = _insert_block_relative(
                    current_anchor, where, block, field_map,
                )


def _block_paragraphs(block, field_map) -> list[tuple[str, bool, Optional[float]]]:
    """Render a DocumentBlock as a list of (text, bold, font_size_pt) tuples."""
    parts: list[tuple[str, bool, Optional[float]]] = []
    title_text = _resolve_block_text(block.title or '', field_map)
    if title_text:
        parts.append((title_text, True, 12.0))
    body_text = _resolve_block_text(block.body or '', field_map)
    for chunk in body_text.split('\n\n'):
        chunk = chunk.strip()
        if chunk:
            parts.append((chunk, False, None))
    return parts


def _append_block_at_end(doc: DocumentType, block, field_map) -> None:
    for text, bold, font_size in _block_paragraphs(block, field_map):
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(text)
        if bold:
            run.bold = True
        if font_size is not None:
            run.font.size = Pt(font_size)


def _insert_block_relative(anchor, where: str, block, field_map):
    """Insert a block's paragraphs relative to ``anchor``, returning the
    last inserted paragraph (so subsequent inserts at the same anchor
    chain forward instead of stacking in reverse).
    """
    parts = _block_paragraphs(block, field_map)
    if not parts:
        return anchor

    last = anchor
    if where == 'after':
        # First paragraph goes right after the anchor; subsequent
        # paragraphs go right after the previously-inserted one.
        for text, bold, font_size in parts:
            last = add_paragraph_after(last, text, bold=bold, font_size_pt=font_size)
        return last
    if where == 'before':
        # First paragraph inserted before the anchor; subsequent ones
        # are inserted before the SAME anchor so order is preserved
        # (they stack above the anchor in the order they were emitted).
        first_inserted = None
        for text, bold, font_size in parts:
            new_p = add_paragraph_before(anchor, text, bold=bold, font_size_pt=font_size)
            if first_inserted is None:
                first_inserted = new_p
        # Next block inserted "before" should land BEFORE this block
        # group, so return the first paragraph we inserted as the new
        # anchor for further `before` operations. For mixed positions
        # this doesn't matter — the caller iterates per-position.
        return first_inserted or anchor
    # 'end' shouldn't reach here (caller short-circuits) but defend.
    _append_block_at_end(anchor._parent, block, field_map)
    return anchor


def _get_uploaded_template(service, template_type) -> Optional[Path]:
    from documents.models import ServiceTemplate
    try:
        template = ServiceTemplate.objects.filter(
            service=service,
            template_type=template_type,
            is_active=True,
        ).first()
        if template and template.file:
            file_path = Path(settings.MEDIA_ROOT) / template.file.name
            if file_path.exists():
                return file_path
    except Exception:
        pass
    return None


def _save_document(doc: DocumentType, prefix: str, request_obj,
                   *, style_tables: bool = True,
                   skip_institutional: bool = False,
                   skip_brand_footer: bool = False,
                   skip_house_style: bool = False) -> str:
    """Persist a generated document, but first run the house-style
    finishing pass so every artifact leaves the pipeline with the same
    typography / colours / footer / institutional header regardless of
    which generator built it. Per-generator code stays focused on the
    *content*; the *form* is centralised here.

    Kwargs let specific generators opt out of pieces of the centralised
    finishing pass:
      * skip_house_style  — don't rewrite fonts/colours/headings.
        Used by the IBTIKAR legacy templates, which ship with their
        own typography already polished by the ESSBO team.
      * skip_institutional — don't inject the DGRSDT banner. Used by
        GENOCLAB-side documents (devis, facture) which have their own
        GENOCLAB letterhead.
      * skip_brand_footer — don't add the "ESSBO — PLAGENOR 4.0 / page
        X / Y" footer.
      * style_tables — repaint every table with the brand header tint.
    """
    # Centralised finishing — these are all idempotent so calling them
    # in addition to whatever the generator already did is safe.
    if not skip_house_style:
        apply_house_style(doc)
    if not skip_institutional:
        ensure_institutional_header(doc)
    if not skip_brand_footer:
        add_brand_footer(doc)
    if style_tables:
        _apply_brand_table_style_everywhere(doc)

    out_dir = Path(settings.MEDIA_ROOT) / 'documents'
    out_dir.mkdir(parents=True, exist_ok=True)
    safe_id = (request_obj.display_id or str(request_obj.pk)).replace('/', '_')
    filename = f"{prefix}_{safe_id}_{datetime.now().strftime('%Y%m%d')}.docx"
    filepath = out_dir / filename
    doc.save(str(filepath))
    return str(filepath)


def _apply_brand_table_style_everywhere(doc: DocumentType) -> None:
    """Apply the unified brand table style to every table in the doc.

    Detects a header row heuristically: if the first row's first cell
    is bold or visibly a label, treat the first row as the header tint;
    otherwise fall back to a left-column accent for the IBTIKAR
    two-column "label : value" layouts. This way every generator picks
    up the same visual without having to refactor each call site.
    """
    for table in doc.tables:
        if not table.rows:
            continue
        # Heuristic: 2-column tables with non-numeric first column are
        # almost always label/value layouts.
        accent = 'header'
        if len(table.columns) == 2 and len(table.rows) >= 2:
            first_col_vals = [
                (r.cells[0].text or '').strip()
                for r in table.rows[:min(4, len(table.rows))]
            ]
            looks_labelled = sum(
                1 for v in first_col_vals if v and not v.replace(' ', '').isdigit()
            ) >= 2
            if looks_labelled:
                accent = 'first-col'
        try:
            style_brand_table(table, accent=accent)
        except Exception:
            # Don't let one mal-formed table break the whole document.
            pass


# Generators ----------------------------------------------------------------

def generate_ibtikar_form(request_obj) -> str:
    """IBTIKAR form. Priority: uploaded template → service-specific
    branded template (egtp_*.docx with the institutional banner already
    in the header) → generic generic template → programmatic fallback.

    Service-specific egtp_*.docx forms ship as printable French forms
    with literal labels ("Nom et prénom : * Nom complet du demandeur")
    instead of ``{{KEY}}`` markers, so they bypass the standard
    substitution pass. The ``apply_legacy_label_substitution`` step fills
    the personal-info / request fields by matching the French label
    patterns and writing the requester's data into the asterisked
    instructional-text slot. Triggered only when the egtp_*.docx path is
    taken — the generic/programmatic paths already use ``{{KEY}}``.
    """
    field_map = build_field_map(request_obj)
    doc: Optional[DocumentType] = None
    using_legacy_form = False

    if request_obj.service:
        uploaded = _get_uploaded_template(request_obj.service, 'IBTIKAR_FORM')
        if uploaded:
            doc = Document(str(uploaded))

    if doc is None and request_obj.service:
        template_name = IBTIKAR_TEMPLATE_MAP.get(request_obj.service.code, '')
        if template_name:
            path = Path(settings.BASE_DIR) / 'documents' / 'docx_templates' / 'ibtikar' / template_name
            if path.exists():
                doc = Document(str(path))
                using_legacy_form = True

    using_generic = False
    if doc is None:
        generic = Path(settings.BASE_DIR) / 'documents' / 'docx_templates' / 'ibtikar_form_template.docx'
        if generic.exists():
            doc = Document(str(generic))
            using_generic = True

    if doc is None:
        doc = _build_ibtikar_form_programmatic(request_obj, field_map)

    replace_placeholders(doc, field_map)
    if using_legacy_form:
        apply_legacy_label_substitution(doc, field_map)
        # Pre-fill the printable sample grid from the digital submission so the
        # lab receives a complete form, not a blank table to re-type by hand.
        populate_legacy_sample_table(doc, request_obj)
        # Fill every "Choisissez un élément" / "Cliquez ici" answer slot from
        # the requester's online service_params (PCR kit, QC level, marker
        # size, recovered volume…). Fuzzy label match → answer; each param
        # used at most once. Without this, the Section 4 questions stayed
        # blank even when the requester answered them online.
        populate_legacy_param_questions(doc, request_obj)
    elif using_generic:
        # The generic template only carries identity placeholders. For a
        # service with no branded egtp form (e.g. one a SuperAdmin created
        # from scratch), render its questions and sample table — with proper
        # human labels — so the generated document carries everything the
        # requester entered.
        _labels = _field_label_map(request_obj)
        _render_service_params(doc, request_obj.service_params, _labels)
        _render_sample_table(doc, request_obj.sample_table, _labels)
    strip_unresolved_placeholders(doc)
    ensure_institutional_header(doc)
    _inject_document_blocks(doc, 'IBTIKAR_FORM', request_obj)
    # The egtp_*.docx IBTIKAR templates are the official printable forms
    # of the platform — they ship with their own typography, table
    # styling and footer. Repainting them with the unified PLAGENOR
    # house style overrides those carefully crafted fonts and colours
    # and gives the result an inconsistent / "loud" feel. So we skip
    # the finishing pass for legacy / generic templates: only the
    # programmatic fallback gets the centralised look.
    if using_legacy_form or using_generic:
        return _save_document(
            doc, 'IBTIKAR', request_obj,
            skip_house_style=True,    # preserve the template's own fonts
            style_tables=False,       # preserve the template's own table look
            skip_brand_footer=True,   # the template already has its footer
        )
    return _save_document(doc, 'IBTIKAR', request_obj)


def _build_ibtikar_form_programmatic(request_obj, field_map) -> DocumentType:
    doc = Document()
    apply_house_style(doc)
    doc.add_heading('Formulaire IBTIKAR — PLAGENOR', level=1)
    doc.add_heading("ESSBO — École Supérieure en Sciences Biologiques d'Oran", level=2)
    doc.add_paragraph(f"Référence : {field_map['DISPLAY_ID']}")
    doc.add_paragraph(f"Date : {field_map['DATE']}")
    doc.add_paragraph('')

    doc.add_heading('Informations du demandeur', level=2)
    table = doc.add_table(rows=8, cols=2)
    table.style = 'Light Grid Accent 1'
    fields = [
        ('Nom complet', field_map['FULL_NAME']),
        ('Email', field_map['EMAIL']),
        ('Téléphone', field_map['PHONE']),
        ('Établissement', field_map['ETABLISSEMENT']),
        ('Laboratoire', field_map['LABORATORY']),
        ('Directeur de recherche', field_map['SUPERVISOR']),
        ('Niveau', field_map['STUDENT_LEVEL']),
        ('Code IBTIKAR-DGRSDT', field_map['IBTIKAR_EXTERNAL_CODE']),
    ]
    for i, (label, value) in enumerate(fields):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = str(value or '')

    doc.add_heading('Service demandé', level=2)
    svc_table = doc.add_table(rows=4, cols=2)
    svc_table.style = 'Light Grid Accent 1'
    svc_fields = [
        ('Code', field_map['SERVICE_CODE']),
        ('Intitulé', field_map['SERVICE_NAME']),
        ('Description', field_map['SERVICE_DESCRIPTION']),
        ('Délai', f"{field_map['SERVICE_TURNAROUND']} jours"),
    ]
    for i, (label, value) in enumerate(svc_fields):
        svc_table.rows[i].cells[0].text = label
        svc_table.rows[i].cells[1].text = str(value or '')

    doc.add_heading('Détails de la demande', level=2)
    doc.add_paragraph(f"Titre du projet : {field_map['TITLE']}")
    if field_map['DESCRIPTION']:
        doc.add_paragraph(f"Description : {field_map['DESCRIPTION']}")
    doc.add_paragraph(f"Urgence : {field_map['URGENCY']}")
    doc.add_paragraph(f"Budget estimé : {field_map['BUDGET_AMOUNT']}")
    doc.add_paragraph(f"Solde IBTIKAR déclaré : {field_map['IBTIKAR_BALANCE']}")

    _render_sample_table(doc, request_obj.sample_table)
    _render_service_params(doc, request_obj.service_params)
    _render_footer(doc)
    return doc


def generate_platform_note(request_obj) -> str:
    """Platform Note. Priority: uploaded template → static generic template
    → programmatic fallback. Substitution is run-preserving so existing
    bold/italic/colour from the template survives.
    """
    field_map = build_field_map(request_obj)
    doc: Optional[DocumentType] = None

    if request_obj.service:
        uploaded = _get_uploaded_template(request_obj.service, 'PLATFORM_NOTE')
        if uploaded:
            doc = Document(str(uploaded))

    if doc is None:
        generic = Path(settings.BASE_DIR) / 'documents' / 'docx_templates' / 'platform_note_template.docx'
        if generic.exists():
            doc = Document(str(generic))

    using_programmatic = doc is None
    if doc is None:
        doc = _build_platform_note_programmatic(request_obj, field_map)

    replace_placeholders(doc, field_map)
    strip_unresolved_placeholders(doc)
    ensure_institutional_header(doc)
    # Tariff justification — appended even when a template was used.
    if not using_programmatic:
        _render_tariff_breakdown(doc, request_obj)
    _inject_document_blocks(doc, 'PLATFORM_NOTE', request_obj)
    # The platform note is meant to be copy-pasted into the DGRSDT
    # IBTIKAR web portal. Word tables paste as tab-separated mush in
    # browser text areas, so we transform any residual table (left over
    # from an uploaded template) into "Label : value" paragraphs first.
    _detableize(doc)
    return _save_document(doc, 'NOTE_PLT', request_obj, style_tables=False)


def _detableize(doc: DocumentType) -> None:
    """Convert every Word table in the document into plain paragraphs.

    Each row is flattened into one paragraph: when the table has two
    columns, the layout is "label : value" with the label bolded; when
    it has more columns, cells are joined with " — ". Empty rows are
    dropped. The table is then removed from the document tree.

    Used by the platform note generator so the resulting DOCX survives
    a Ctrl+C / Ctrl+V into the DGRSDT portal's text areas, where Word
    tables paste as tabs / line breaks chaos.
    """
    for table in list(doc.tables):
        anchor = table._element  # the <w:tbl> xml element
        parent = anchor.getparent()
        if parent is None:
            continue
        idx = list(parent).index(anchor)

        # Build the replacement paragraphs first (before removing the
        # table so we still have access to its cells).
        new_paragraphs = []
        ncols = len(table.columns)
        for row in table.rows:
            cells = [(c.text or '').strip() for c in row.cells]
            if not any(cells):
                continue
            if ncols == 2:
                label, value = cells[0], cells[1]
                p = doc.add_paragraph()
                rl = p.add_run(f"{label} : " if label else '')
                rl.bold = True
                p.add_run(value)
            else:
                text = ' — '.join(c for c in cells if c)
                p = doc.add_paragraph(text)
            # add_paragraph appends to body; move it in front of the table.
            new_paragraphs.append(p._element)

        # Move the freshly-added paragraphs to the table's position, then
        # drop the table.
        for offset, new_p in enumerate(new_paragraphs):
            new_p.getparent().remove(new_p)
            parent.insert(idx + offset, new_p)
        parent.remove(anchor)


def _summarise_samples(samples, total_n) -> str:
    """Return a one-line sample summary for the platform note.

    The DGRSDT IBTIKAR portal doesn't need the detailed per-sample grid
    (that lives in the lab's reception form). The platform note only
    needs to know *how many* samples and, when an "isolat type" column
    is heterogeneous, the count per type.

    Strategy:
      * Look for the first categorical column that varies across rows
        (organism_type, sample_origin, dna_type, …).
      * If two or more distinct values exist, format as
        "N (X bactéries, Y levures, …)".
      * Otherwise return just "N échantillon(s)".
    """
    if not samples or total_n <= 0:
        return f"{total_n}" if total_n else "—"
    # Common categorical column names — keep the order: most specific first.
    candidate_cols = ('organism_type', 'sample_origin', 'dna_type',
                      'primer_type', 'sequencing_mode', 'source')
    chosen = None
    for col in candidate_cols:
        values = [str((s or {}).get(col, '')).strip() for s in samples if isinstance(s, dict)]
        values = [v for v in values if v]
        if len(set(values)) >= 2:
            chosen = (col, values)
            break
    base = f"{total_n} échantillon" + ('s' if total_n > 1 else '')
    if chosen is None:
        return base
    _, values = chosen
    from collections import Counter
    counts = Counter(values)
    parts = [f"{c} {label}" for label, c in counts.most_common()]
    return f"{base} ({', '.join(parts)})"


def _render_tariff_breakdown(doc, request_obj) -> None:
    """Render an itemised tariff justification in the platform note.

    Renders a styled section that breaks the total down by:
      * Type d'analyse (service code + name)
      * Mode d'analyse / niveau retenu (analysis_mode / qc_level / …)
      * Caractère pathogène (Oui / Non)
      * Prix de base unitaire
      * Multiplicateur appliqué
      * Nombre d'échantillons (résumé : compte total + ventilation par
        type quand l'isolat est hétérogène — pas le détail ligne à ligne)
      * Total = base × multiplicateur × N (utilise admin_validated_price
        si l'admin_ops a réajusté le tarif)

    Source of truth: re-runs ``resolve_cost`` from ``core.pricing`` so the
    figures here are the SAME ones the requester saw live in the form.
    Defensive against missing data — every label degrades to "—".
    """
    from core.pricing import resolve_cost

    doc.add_heading('Justification du tarif', level=2)

    params = request_obj.service_params or {}
    samples = request_obj.sample_table or []
    n = len([s for s in samples if isinstance(s, dict) and any(v not in (None, '', [], {}) for v in s.values())])
    if n == 0:
        n = len(samples) or 1

    result = {}
    if request_obj.service:
        try:
            result = resolve_cost(
                request_obj.service, request_obj.channel,
                sample_table=samples, service_params=params,
                urgency=request_obj.urgency or 'Normal',
            )
        except Exception:
            result = {}

    breakdown = result.get('breakdown', {}) or {}
    # resolve_cost returns a dict for the YAML/programmatic paths and a
    # list for the DB-tier path. Accept either and back-fill the missing
    # fields directly from the service_params + YAML when we got a list.
    if isinstance(breakdown, list):
        breakdown = {}

    # First pass: lift whatever the dict-breakdown gave us.
    base_price = breakdown.get('base_price')
    multiplier = breakdown.get('multiplier')
    mult_key = breakdown.get('multiplier_key')
    pathogenic = breakdown.get('pathogenic')

    # Second pass: fall back to the service definition + params when the
    # cost path didn't provide a dict breakdown (DB tiers, OVERRIDE, …).
    if pathogenic is None and 'pathogenic' in params:
        pathogenic = bool(params.get('pathogenic'))
    if mult_key is None:
        for k in ('analysis_mode', 'qc_level', 'sequencing_mode',
                  'drying_level', 'primer_type'):
            v = params.get(k)
            if v:
                mult_key = str(v)
                break
    # If still missing, dig into the YAML registry for base price /
    # multiplier so the four target fields never stay at "—" on a
    # well-priced service.
    if base_price is None or multiplier is None:
        try:
            from core.registry import get_service_def
            sdef = get_service_def(getattr(request_obj.service, 'code', '')) or {}
            pricing = sdef.get('pricing', {}) or {}
            bp_map = pricing.get('base_price', {}) or {}
            mult_map = pricing.get('multipliers', {}) or {}
            if base_price is None:
                key = 'pathogenic' if pathogenic else 'non_pathogenic'
                base_price = bp_map.get(key) or bp_map.get('default')
            if multiplier is None and mult_key is not None:
                multiplier = mult_map.get(str(mult_key))
                if multiplier is None:
                    multiplier = mult_map.get(mult_key)
        except Exception:
            pass

    # Total — prefer admin_validated_price (admin re-pricing post-
    # submission) over the live resolver result, over the saved
    # budget_amount.
    admin_price = getattr(request_obj, 'admin_validated_price', None)
    if admin_price:
        total = float(admin_price)
        price_source = 'Réajusté par administration'
    else:
        total = result.get('total') or request_obj.budget_amount or 0
        price_source = None

    # Sample summary — count total + breakdown by primary categorical
    # column (organism_type, etc.) when the column is heterogeneous.
    samples_label = _summarise_samples(samples, n)

    rows = [
        ("Type d'analyse",
         f"{getattr(request_obj.service, 'code', '—')} — {getattr(request_obj.service, 'name', '')}"),
        ("Mode / niveau retenu", str(mult_key) if mult_key else '—'),
        ("Caractère pathogène",
         '—' if pathogenic is None else ('Oui' if pathogenic else 'Non')),
        ("Prix de base unitaire",
         f"{float(base_price):,.0f} DZD".replace(',', ' ')
         if isinstance(base_price, (int, float)) else '—'),
        ("Multiplicateur appliqué",
         f"× {multiplier}" if multiplier not in (None, '—') else '—'),
        ("Nombre d'échantillons", samples_label),
        ("Total",
         f"{float(total):,.0f} DZD".replace(',', ' ')
         if isinstance(total, (int, float)) else str(total)),
    ]
    if price_source:
        rows.append(("Source du tarif", price_source))
    # Rendered as plain "Label : value" paragraphs (no Word table). The
    # platform note is meant to be copy-pasted into the DGRSDT IBTIKAR
    # portal, where pasted Word tables turn into tab-separated mush in
    # the textareas. Paragraphs survive copy-paste cleanly.
    for label, value in rows:
        p = doc.add_paragraph()
        run_l = p.add_run(f"{label} : ")
        run_l.bold = True
        run_v = p.add_run(str(value))
        # Make the Total line visually stand out without using a table.
        if label == 'Total':
            run_l.font.size = Pt(12)
            run_v.font.size = Pt(12)
            run_v.bold = True

    doc.add_paragraph(
        "Formule appliquée : Prix de base × Multiplicateur × Nombre d'échantillons.",
        style='Intense Quote' if 'Intense Quote' in [s.name for s in doc.styles] else None,
    )


def _build_platform_note_programmatic(request_obj, field_map) -> DocumentType:
    doc = Document()
    apply_house_style(doc)

    doc.add_heading('NOTE DE PLATEFORME — PLAGENOR', level=1)
    doc.add_paragraph("ESSBO — École Supérieure en Sciences Biologiques d'Oran")
    doc.add_paragraph(f"Référence : {field_map['DISPLAY_ID']}")
    doc.add_paragraph(f"Date d'émission : {field_map['DATETIME']}")
    doc.add_paragraph('')

    doc.add_heading('Demandeur', level=2)
    # Plain "Label : value" paragraphs — no Word table. The note is
    # designed to be copy-pasted into the DGRSDT IBTIKAR portal text
    # areas, where Word tables become tab-separated garbage on paste.
    fields = [
        ('Nom complet', field_map['FULL_NAME']),
        ('Établissement', field_map['ETABLISSEMENT']),
        ('Laboratoire', field_map['LABORATORY']),
        ('Niveau / fonction', field_map['STUDENT_LEVEL']),
        ('Directeur de recherche', field_map['SUPERVISOR']),
        ('Email', field_map['EMAIL']),
        ('Téléphone', field_map['PHONE']),
    ]
    for label, value in fields:
        p = doc.add_paragraph()
        run_l = p.add_run(f"{label} : ")
        run_l.bold = True
        p.add_run(str(value or ''))

    doc.add_heading('Service demandé', level=2)
    doc.add_paragraph(f"Code : {field_map['SERVICE_CODE']}")
    doc.add_paragraph(f"Intitulé : {field_map['SERVICE_NAME']}")
    if field_map['SERVICE_DESCRIPTION']:
        doc.add_paragraph(f"Description : {field_map['SERVICE_DESCRIPTION']}")
    doc.add_paragraph(f"Délai estimé : {field_map['SERVICE_TURNAROUND']} jours ouvrables")

    doc.add_heading('Détails de la demande', level=2)
    doc.add_paragraph(f"Titre : {field_map['TITLE']}")
    if field_map['DESCRIPTION']:
        doc.add_paragraph(f"Description : {field_map['DESCRIPTION']}")
    doc.add_paragraph(f"Canal : {field_map['CHANNEL']}")
    doc.add_paragraph(f"Urgence : {field_map['URGENCY']}")

    _render_service_params(doc, request_obj.service_params)
    _render_sample_table(doc, request_obj.sample_table)

    doc.add_heading('Décompte budgétaire IBTIKAR', level=2)
    doc.add_paragraph('Budget annuel par étudiant : 200 000 DZD')
    doc.add_paragraph(f"Montant de cette prestation : {field_map['BUDGET_AMOUNT']}")
    if request_obj.declared_ibtikar_balance:
        doc.add_paragraph(f"Solde IBTIKAR déclaré : {field_map['IBTIKAR_BALANCE']}")

    # Tariff justification — see _render_tariff_breakdown docstring.
    _render_tariff_breakdown(doc, request_obj)

    if request_obj.assigned_to:
        doc.add_heading('Assignation', level=2)
        doc.add_paragraph(f"Analyste : {field_map['ASSIGNED_ANALYST']}")
        if field_map['ANALYST_EMAIL']:
            doc.add_paragraph(f"Email analyste : {field_map['ANALYST_EMAIL']}")
        if field_map['APPOINTMENT_DATE'] != 'Non défini':
            doc.add_paragraph(f"Rendez-vous : {field_map['APPOINTMENT_DATE']}")

    _render_footer(doc)
    return doc


def generate_quote(request_obj) -> str:
    """GENOCLAB quote (Facture Proforma).

    Renders the SAIDAL-style commercial layout: GENOCELAB logo,
    issuer + client header (issuer details are CMS-editable so the
    SuperAdmin can refresh bank accounts or NIF without code), table
    of prestations × prix × montant, HT / TVA / TTC totals, and a
    legal footer block. No ESSBO/PLAGENOR institutional banner — this
    is the commercial side of the platform, not academic.
    """
    field_map = build_field_map(request_obj)

    # Quote number — sequential, atomic.
    from core.sequences import next_value
    year = datetime.now().year
    seq = next_value(f'GENOCLAB-QUOTE-{year}')
    quote_number = f"GENOCLAB-DEV-{year}-{seq:04d}"
    field_map['QUOTE_NUMBER'] = quote_number
    field_map['INVOICE_NUMBER'] = quote_number  # legacy alias

    doc = _build_genoclab_doc(
        title_key='genoclab_quote_title',
        doc_number=quote_number,
        request_obj=request_obj,
        line_items=(request_obj.quote_detail or {}).get('items') or [],
        admin_fees=(request_obj.quote_detail or {}).get('admin_fees', 0),
        report_fees=(request_obj.quote_detail or {}).get('report_fees', 0),
        vat_rate=(request_obj.quote_detail or {}).get('vat_rate'),
    )
    _inject_document_blocks(doc, 'QUOTE', request_obj)
    # Skip the ESSBO institutional header + ESSBO/PLAGENOR footer —
    # GENOCLAB documents have their own header/footer built above.
    return _save_document(
        doc, 'DEVIS', request_obj,
        style_tables=False,        # tables are already styled by our layout
        skip_institutional=True,   # GENOCELAB header instead
        skip_brand_footer=True,    # GENOCLAB-specific footer text
    )


def _build_genoclab_doc(
    *, title_key: str, doc_number: str, request_obj,
    line_items, admin_fees=0, report_fees=0, vat_rate=None,
) -> DocumentType:
    """Common builder for the GENOCLAB quote and invoice — assembles
    header + prestation table + totals + legal footer in one place so
    quote and invoice always carry the same look and field discipline.
    """
    from documents.genoclab_layout import (
        add_genoclab_header, add_prestation_table, add_genoclab_footer,
        cms_get,
    )
    doc = Document()
    apply_house_style(doc)

    # Client coords — pulled from the requester profile, with sensible
    # blanks when something is missing.
    requester = getattr(request_obj, 'requester', None)
    client_name = ''
    client_lines = []
    if requester is not None:
        client_name = requester.get_full_name() or requester.username or ''
        org = getattr(requester, 'organization', '')
        phone = getattr(requester, 'phone', '')
        email = getattr(requester, 'email', '')
        if org: client_lines.append(org)
        if phone: client_lines.append(f"Tél : {phone}")
        if email: client_lines.append(email)

    add_genoclab_header(
        doc,
        title=cms_get(title_key),
        doc_number=doc_number,
        doc_date=datetime.now().strftime('%d/%m/%Y'),
        client_name=client_name,
        client_lines=client_lines,
    )

    # Add a small spacer paragraph.
    doc.add_paragraph()

    # Build the prestation list. admin_fees / report_fees come from the
    # quote_detail and become their own line items so they're visible in
    # the table — the SuperAdmin / client sees exactly where every DA
    # goes.
    items = list(line_items or [])
    extras = []
    if admin_fees and float(admin_fees) > 0:
        extras.append({'label': 'Frais administratifs',
                       'quantity': 1, 'unit_price': float(admin_fees),
                       'total': float(admin_fees)})
    if report_fees and float(report_fees) > 0:
        extras.append({'label': 'Frais de rapport',
                       'quantity': 1, 'unit_price': float(report_fees),
                       'total': float(report_fees)})

    grand_total = add_prestation_table(doc, items + extras, vat_rate=vat_rate)
    add_genoclab_footer(doc, total_amount=grand_total)
    return doc


def _build_quote_programmatic(request_obj, field_map) -> DocumentType:
    doc = Document()
    apply_house_style(doc)

    doc.add_heading('DEVIS — GENOCLAB', level=1)
    doc.add_paragraph("ESSBO — École Supérieure en Sciences Biologiques d'Oran")
    doc.add_paragraph(f"N° Devis : {field_map['QUOTE_NUMBER']}")
    doc.add_paragraph(f"Référence demande : {field_map['DISPLAY_ID']}")
    doc.add_paragraph(f"Date : {field_map['DATE']}")
    doc.add_paragraph('')

    doc.add_heading('Client', level=2)
    client_table = doc.add_table(rows=5, cols=2)
    client_table.style = 'Light Grid Accent 1'
    fields = [
        ('Nom', field_map['CLIENT_NAME']),
        ('Organisation', field_map['ORGANIZATION']),
        ('Email', field_map['CLIENT_EMAIL']),
        ('Téléphone', field_map['PHONE']),
        ('Laboratoire', field_map['LABORATORY']),
    ]
    for i, (label, value) in enumerate(fields):
        client_table.rows[i].cells[0].text = label
        client_table.rows[i].cells[1].text = str(value or '')

    doc.add_heading('Prestations', level=2)
    table = doc.add_table(rows=2, cols=4)
    table.style = 'Light Grid Accent 1'
    headers = ['Description', 'Quantité', 'Prix unitaire', 'Total']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    table.rows[1].cells[0].text = field_map['SERVICE_NAME']
    table.rows[1].cells[1].text = '1'
    table.rows[1].cells[2].text = field_map['SUBTOTAL_HT']
    table.rows[1].cells[3].text = field_map['SUBTOTAL_HT']

    doc.add_paragraph('')
    summary = doc.add_table(rows=3, cols=2)
    summary.style = 'Light Grid Accent 1'
    summary.rows[0].cells[0].text = 'Sous-total HT'
    summary.rows[0].cells[1].text = field_map['SUBTOTAL_HT']
    summary.rows[1].cells[0].text = f"TVA ({field_map['VAT_RATE']})"
    summary.rows[1].cells[1].text = field_map['VAT_AMOUNT']
    summary.rows[2].cells[0].text = 'Total TTC'
    summary.rows[2].cells[1].text = field_map['TOTAL_TTC']

    _render_footer(doc)
    return doc


def generate_reception_form(request_obj) -> str:
    field_map = build_field_map(request_obj)
    doc: Optional[DocumentType] = None

    if request_obj.service:
        uploaded = _get_uploaded_template(request_obj.service, 'RECEPTION_FORM')
        if uploaded:
            doc = Document(str(uploaded))

    if doc is None:
        generic = Path(settings.BASE_DIR) / 'documents' / 'docx_templates' / 'reception_form_template.docx'
        if generic.exists():
            doc = Document(str(generic))

    if doc is None:
        doc = _build_reception_form_programmatic(request_obj, field_map)

    replace_placeholders(doc, field_map)
    strip_unresolved_placeholders(doc)
    ensure_institutional_header(doc)
    _inject_document_blocks(doc, 'RECEPTION_FORM', request_obj)
    return _save_document(doc, 'RECEPTION', request_obj)


def _build_reception_form_programmatic(request_obj, field_map) -> DocumentType:
    doc = Document()
    apply_house_style(doc)

    doc.add_heading("Fiche de Réception d'Échantillons", level=1)
    doc.add_paragraph('PLAGENOR — ESSBO')
    doc.add_paragraph(f"Référence : {field_map['DISPLAY_ID']}")
    doc.add_paragraph(f"Code de suivi : {field_map['TRACKING_CODE']}")
    doc.add_paragraph('')

    table = doc.add_table(rows=6, cols=2)
    table.style = 'Light Grid Accent 1'
    fields = [
        ('Service', field_map['SERVICE_NAME']),
        ('Canal', field_map['CHANNEL']),
        ('Urgence', field_map['URGENCY']),
        ('Date de RDV', field_map['APPOINTMENT_DATE']),
        ('Analyste assigné', field_map['ASSIGNED_ANALYST']),
        ('Date de soumission', field_map['SUBMISSION_DATE']),
    ]
    for i, (label, value) in enumerate(fields):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = str(value or '')

    doc.add_heading('Déposant', level=2)
    client_table = doc.add_table(rows=5, cols=2)
    client_table.style = 'Light Grid Accent 1'
    client_fields = [
        ('Nom', field_map['FULL_NAME']),
        ('Email', field_map['EMAIL']),
        ('Téléphone', field_map['PHONE']),
        ('Établissement', field_map['ETABLISSEMENT']),
        ('Laboratoire', field_map['LABORATORY']),
    ]
    for i, (label, value) in enumerate(client_fields):
        client_table.rows[i].cells[0].text = label
        client_table.rows[i].cells[1].text = str(value or '')

    _render_sample_table(doc, request_obj.sample_table)

    doc.add_heading('Réception', level=2)
    rec_table = doc.add_table(rows=4, cols=2)
    rec_table.style = 'Light Grid Accent 1'
    rec_fields = [
        ('Date de réception', '___ / ___ / ______'),
        ("Nombre d'échantillons reçus", '____________'),
        ('État des échantillons', '☐ Bon   ☐ Acceptable   ☐ Dégradé'),
        ('Observations', ''),
    ]
    for i, (label, value) in enumerate(rec_fields):
        rec_table.rows[i].cells[0].text = label
        rec_table.rows[i].cells[1].text = str(value or '')

    doc.add_paragraph('')
    doc.add_paragraph('Signature du réceptionniste : ________________________')
    doc.add_paragraph('Signature du déposant : ________________________')
    _render_footer(doc)
    return doc


# Section helpers (shared body builders) -----------------------------------

def _field_label_map(request_obj) -> dict:
    """Map field/column ``name`` → human label for the request's service.

    Pulls from the YAML registry (parameters + sample_table columns, FR label
    preferred) and the DB ``custom_fields`` a SuperAdmin defined, so the
    programmatic document shows "Objectif de l'analyse" instead of the raw
    ``analysis_goal`` key.
    """
    labels = {}
    svc = getattr(request_obj, 'service', None)
    code = getattr(svc, 'code', '') or ''
    try:
        from core.registry import get_service_def
        sdef = get_service_def(code) or {}
        for p in sdef.get('parameters', []) or []:
            if p.get('name'):
                labels[p['name']] = p.get('label_fr') or p.get('label') or p['name']
        for c in (sdef.get('sample_table', {}) or {}).get('columns', []) or []:
            if c.get('name'):
                labels[c['name']] = c.get('label') or c['name']
    except Exception:
        pass
    try:
        if svc is not None:
            for f in svc.custom_fields.all():
                if f.name:
                    labels[f.name] = f.label or f.name
    except Exception:
        pass
    return labels


def _render_sample_table(doc: DocumentType, sample_table, label_map=None) -> None:
    if not sample_table or not isinstance(sample_table, list):
        return
    samples = [s for s in sample_table if isinstance(s, dict) and any(s.values())]
    if not samples:
        return
    label_map = label_map or {}
    doc.add_heading('Tableau des échantillons', level=3)
    headers = list(samples[0].keys())
    table = doc.add_table(rows=len(samples) + 1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    for j, h in enumerate(headers):
        table.rows[0].cells[j].text = label_map.get(h) or h.replace('_', ' ').capitalize()
    for i, sample in enumerate(samples):
        for j, h in enumerate(headers):
            table.rows[i + 1].cells[j].text = str(sample.get(h, ''))


def _render_service_params(doc: DocumentType, service_params, label_map=None) -> None:
    if not service_params or not isinstance(service_params, dict):
        return
    non_empty = [(k, v) for k, v in service_params.items() if v not in (None, '', [], {})]
    if not non_empty:
        return
    label_map = label_map or {}
    doc.add_heading('Paramètres du service', level=3)
    table = doc.add_table(rows=len(non_empty), cols=2)
    table.style = 'Light Grid Accent 1'
    for i, (key, value) in enumerate(non_empty):
        clean = key.replace('param_', '')
        table.rows[i].cells[0].text = label_map.get(clean) or clean.replace('_', ' ').capitalize()
        table.rows[i].cells[1].text = str(value)


def _render_footer(doc: DocumentType) -> None:
    doc.add_paragraph('')
    doc.add_paragraph('—' * 40)
    p = doc.add_paragraph()
    run = p.add_run(
        f"Document généré automatiquement par PLAGENOR 4.0 · "
        f"{datetime.now().strftime('%d/%m/%Y à %H:%M')}"
    )
    run.font.size = Pt(9)


def generate_stats_report(bundle: dict, filters: dict, actor) -> str:
    """Generate the official statistics report DOCX.

    Branded like every other PLAGENOR document — institutional header is
    injected by ``ensure_institutional_header`` — so the PDF conversion
    that follows yields a presentation-ready report.
    """
    doc = Document()
    apply_house_style(doc)

    doc.add_heading('PLAGENOR 4.0 — Statistiques institutionnelles', level=1)
    doc.add_paragraph(
        f"ESSBO — École Supérieure en Sciences Biologiques d'Oran")
    doc.add_paragraph(
        f"Édité le {datetime.now().strftime('%d/%m/%Y à %H:%M')} "
        f"par {actor.get_full_name() or actor.username}"
    )

    # Active filters block
    if filters:
        doc.add_heading('Filtres appliqués', level=3)
        label_map = {
            'date_from': 'Du', 'date_to': 'Au', 'channel': 'Canal',
            'service_code': 'Service', 'status': 'Statut',
            'wilaya': 'Wilaya', 'organization': 'Établissement',
            'gender': 'Sexe', 'analysis_frame': "Cadre d'analyse",
        }
        for k, v in filters.items():
            doc.add_paragraph(f"• {label_map.get(k, k)} : {v}")

    # Headline KPIs
    kpis = bundle.get('kpis', {})
    doc.add_heading('Indicateurs principaux', level=2)
    kpi_table = doc.add_table(rows=9, cols=2)
    kpi_table.style = 'Light Grid Accent 1'
    rows = [
        ('Demandes', kpis.get('total', 0)),
        ('Complétées', kpis.get('completed', 0)),
        ('En cours', kpis.get('in_progress', 0)),
        ('Rejetées', kpis.get('rejected', 0)),
        ('Taux de complétion', f"{kpis.get('completion_rate', 0)} %"),
        ('Demandes IBTIKAR', kpis.get('ibtikar_count', 0)),
        ('Demandes GENOCLAB', kpis.get('genoclab_count', 0)),
        ('Revenu virtuel IBTIKAR', f"{kpis.get('ibtikar_virtual_revenue', 0):,.0f} DA"),
        ('Revenu GENOCLAB', f"{kpis.get('genoclab_revenue', 0):,.0f} DA"),
    ]
    for i, (label, value) in enumerate(rows):
        kpi_table.rows[i].cells[0].text = label
        kpi_table.rows[i].cells[1].text = str(value)

    def _section(title, key, col1='Catégorie'):
        data = bundle.get(key)
        if not data:
            return
        doc.add_heading(title, level=2)
        t = doc.add_table(rows=len(data) + 1, cols=2)
        t.style = 'Light Grid Accent 1'
        t.rows[0].cells[0].text = col1
        t.rows[0].cells[1].text = 'Demandes'
        for i, r in enumerate(data, start=1):
            t.rows[i].cells[0].text = str(r.get('label', '—'))
            t.rows[i].cells[1].text = str(r.get('count', 0))

    _section('Répartition par service', 'by_service', col1='Service')
    _section('Répartition par statut', 'by_status', col1='Statut')
    _section('Répartition par wilaya', 'by_wilaya', col1='Wilaya')
    _section('Répartition par établissement', 'by_organization', col1='Établissement')
    _section("Répartition par cadre d'analyse", 'by_analysis_frame', col1='Cadre')
    _section('Répartition par sexe', 'by_gender', col1='Sexe')

    trend = bundle.get('trend') or []
    if trend:
        doc.add_heading('Tendance mensuelle', level=2)
        t = doc.add_table(rows=len(trend) + 1, cols=2)
        t.style = 'Light Grid Accent 1'
        t.rows[0].cells[0].text = 'Mois'
        t.rows[0].cells[1].text = 'Demandes'
        for i, r in enumerate(trend, start=1):
            t.rows[i].cells[0].text = r['month']
            t.rows[i].cells[1].text = str(r['count'])

    _render_footer(doc)
    ensure_institutional_header(doc)

    out_dir = Path(settings.MEDIA_ROOT) / 'documents'
    out_dir.mkdir(parents=True, exist_ok=True)
    filename = f"PLAGENOR_Statistiques_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = out_dir / filename
    doc.save(str(filepath))
    return str(filepath)


def generate_invoice_document(invoice_obj) -> str:
    """GENOCLAB invoice (final, post-payment).

    Same SAIDAL-style layout as the quote (generate_quote): GENOCELAB
    logo, CMS-editable issuer block, client block, prestation grid,
    HT / TVA / TTC totals, legal footer. Differs only in the title
    (Facture, from the genoclab_invoice_title CMS key) and the source
    of the line items (the Invoice's line_items JSON).
    """
    from documents.genoclab_layout import (
        add_genoclab_header, add_prestation_table, add_genoclab_footer,
        cms_get,
    )
    doc = Document()
    apply_house_style(doc)

    # Client coords from the invoice row.
    client = invoice_obj.client
    client_name = ''
    client_lines = []
    if client is not None:
        client_name = client.get_full_name() or client.username or ''
        org = getattr(client, 'organization', '')
        phone = getattr(client, 'phone', '')
        email = getattr(client, 'email', '')
        if org: client_lines.append(org)
        if phone: client_lines.append(f"Tél : {phone}")
        if email: client_lines.append(email)
    elif invoice_obj.request and invoice_obj.request.requester:
        client_name = invoice_obj.request.requester.get_full_name() or ''

    add_genoclab_header(
        doc,
        title=cms_get('genoclab_invoice_title'),
        doc_number=invoice_obj.invoice_number,
        doc_date=invoice_obj.created_at.strftime('%d/%m/%Y'),
        client_name=client_name,
        client_lines=client_lines,
    )
    doc.add_paragraph()

    # line_items on Invoice carries 'description' instead of 'label'.
    items = []
    for it in (invoice_obj.line_items or []):
        items.append({
            'label': it.get('description') or it.get('label') or '',
            'quantity': it.get('quantity', 0),
            'unit_price': it.get('unit_price', 0),
            'total': it.get('total'),
        })

    grand_total = add_prestation_table(doc, items, vat_rate=float(invoice_obj.vat_rate or 0.19))
    add_genoclab_footer(doc, total_amount=grand_total)

    out_dir = Path(settings.MEDIA_ROOT) / 'documents'
    out_dir.mkdir(parents=True, exist_ok=True)
    filename = f"FACTURE_{invoice_obj.invoice_number}.docx"
    filepath = out_dir / filename
    doc.save(str(filepath))
    return str(filepath)
