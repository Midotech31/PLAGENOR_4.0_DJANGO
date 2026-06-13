"""DOCX helpers: run-preserving placeholder substitution + header banner injection.

Phase 3.7 extracted these out of generators.py so all four document generators
share a single, correct implementation. The previous approach in generators
(``paragraph.text = paragraph.text.replace(...)``) silently collapsed every
formatted run into a single un-styled run, destroying bold/italic/font from
the source template. The substitution here walks run boundaries so styling
survives. The header-banner helper writes the institutional logo into
``word/header*.xml`` so Platform Note / Quote / Reception render with the
same letterhead that the IBTIKAR templates already carry.
"""
from __future__ import annotations

import re
from pathlib import Path
from typing import Mapping, Optional

from docx import Document
from docx.document import Document as DocumentType
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

# Re-exports for callers: positional insertion primitives live in this
# module so the generators only need a single import statement. See
# ``add_paragraph_after`` / ``_find_anchor_for_position`` further down.


_PLACEHOLDER_RE = re.compile(r'\{\{[A-Z0-9_]+\}\}')


def _substitute_in_runs(paragraph, replacements: Mapping[str, str]) -> None:
    """Replace ``{{KEY}}`` placeholders in a paragraph while preserving runs.

    Strategy: rebuild the joined text, apply replacements, then write the
    result into the first run and blank the rest. Cross-run placeholders
    (Word likes to split tokens across runs when spell-check or language
    tags fire) are handled because we always operate on the joined string.
    The first run's formatting wins for the entire paragraph after a hit;
    this is the same trade-off ``python-docx-template`` makes and is the
    only safe choice without sophisticated run-by-run tokenisation.
    """
    runs = paragraph.runs
    if not runs:
        return
    joined = ''.join(r.text or '' for r in runs)
    if not joined:
        return

    new_text = joined
    for key, value in replacements.items():
        if key in new_text:
            new_text = new_text.replace(key, str(value if value is not None else ''))

    if new_text == joined:
        return

    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ''


def replace_placeholders(doc: DocumentType, field_map: Mapping[str, str]) -> None:
    """Replace ``{{KEY}}`` placeholders everywhere in the document.

    ``field_map`` maps bare keys (``'FULL_NAME'``) to their string values;
    this function wraps them in ``{{...}}`` braces and applies them to:
    paragraphs (body), table cells (recursively, for nested tables), and
    header/footer paragraphs in every section. Bare-key replacement is
    intentionally NOT supported here — it caused the previous corruption
    where ``replace('APPOINTMENT_DATE', '24/05/2026')`` rewrote
    ``{{APPOINTMENT_DATE}}`` into ``{{APPOINTMENT_24/05/2026}}``.
    """
    braced: dict[str, str] = {
        f'{{{{{key}}}}}': '' if value is None else str(value)
        for key, value in field_map.items()
    }

    def visit_paragraphs(paragraphs):
        for p in paragraphs:
            _substitute_in_runs(p, braced)

    def visit_tables(tables):
        for t in tables:
            for row in t.rows:
                for cell in row.cells:
                    visit_paragraphs(cell.paragraphs)
                    visit_tables(cell.tables)

    visit_paragraphs(doc.paragraphs)
    visit_tables(doc.tables)

    for section in doc.sections:
        for header in (section.header, section.first_page_header, section.even_page_header):
            if header is not None:
                visit_paragraphs(header.paragraphs)
                visit_tables(header.tables)
        for footer in (section.footer, section.first_page_footer, section.even_page_footer):
            if footer is not None:
                visit_paragraphs(footer.paragraphs)
                visit_tables(footer.tables)


def strip_unresolved_placeholders(doc: DocumentType) -> None:
    """Remove any ``{{...}}`` markers still present after substitution.

    A stray placeholder reaching the user is worse than a missing value;
    this helper sweeps after ``replace_placeholders`` to keep the output
    clean even when a template uses a key the field_map does not provide.
    """
    def visit(paragraphs):
        for p in paragraphs:
            joined = ''.join(r.text or '' for r in p.runs)
            if not joined or not _PLACEHOLDER_RE.search(joined):
                continue
            cleaned = _PLACEHOLDER_RE.sub('', joined)
            if p.runs:
                p.runs[0].text = cleaned
                for r in p.runs[1:]:
                    r.text = ''

    def visit_tables(tables):
        for t in tables:
            for row in t.rows:
                for cell in row.cells:
                    visit(cell.paragraphs)
                    visit_tables(cell.tables)

    visit(doc.paragraphs)
    visit_tables(doc.tables)
    for section in doc.sections:
        for header in (section.header, section.first_page_header, section.even_page_header):
            if header is not None:
                visit(header.paragraphs)
                visit_tables(header.tables)
        for footer in (section.footer, section.first_page_footer, section.even_page_footer):
            if footer is not None:
                visit(footer.paragraphs)
                visit_tables(footer.tables)


# Banner / letterhead -------------------------------------------------------

_LOGO_BANNER = Path(__file__).resolve().parent / 'assets' / 'institutional_banner.png'


def ensure_institutional_header(doc: DocumentType, banner_path: Optional[Path] = None) -> None:
    """Inject the institutional logo banner into the document's primary header.

    Skips silently if the document's header already contains an image
    (avoids double-stamping templates that ship with their own banner —
    e.g. the egtp_*.docx IBTIKAR forms). Skips silently if the PNG asset
    is missing so a missing-file in dev never breaks a generation.
    """
    banner = banner_path or _LOGO_BANNER
    if not banner.exists():
        return
    for section in doc.sections:
        header = section.header
        if header is None:
            continue
        has_image = any(
            run.element.findall(qn('w:drawing'))
            for paragraph in header.paragraphs
            for run in paragraph.runs
        )
        if has_image:
            continue
        if header.paragraphs:
            paragraph = header.paragraphs[0]
            paragraph.text = ''
        else:
            paragraph = header.add_paragraph()
        run = paragraph.add_run()
        try:
            run.add_picture(str(banner), width=Inches(6.5))
        except Exception:
            return


def apply_house_style(doc: DocumentType) -> None:
    """Apply the PLAGENOR house style to a programmatically built document.

    A4 page size, symmetric 1" margins, body font 11pt. Idempotent — running
    twice on the same doc produces the same result.
    """
    A4_WIDTH = Inches(8.27)
    A4_HEIGHT = Inches(11.69)
    for section in doc.sections:
        section.page_width = A4_WIDTH
        section.page_height = A4_HEIGHT
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
    try:
        style = doc.styles['Normal']
        style.font.size = Pt(11)
    except (KeyError, AttributeError):
        pass


# Positional paragraph insertion -------------------------------------------

# Anchor matching: case-insensitive, whitespace-collapsed substring match
# against the visible text of a heading-shaped paragraph. A paragraph counts
# as "heading-shaped" when its style is Heading[N], OR its leading text
# matches "<digit>." (the egtp_*.docx IBTIKAR templates use plain-text
# numbered titles like "1. Informations du demandeur" instead of Heading
# styles), OR one of its runs has font size >= 14pt and is bold.
_NUMBERED_HEADING_RE = re.compile(r'^\s*\d+[.)]\s')

_ANCHOR_KEYWORDS = {
    'AFTER_REQUESTER': (
        'demandeur', 'déposant', 'client', 'requester', 'requérant',
        'المُقدم', 'الطالب', 'العميل',
    ),
    'AFTER_SAMPLES': (
        'échantillon', 'samples', 'sample', 'tableau des échantillons',
        'العينات', 'عينة',
    ),
}


def _is_heading_shaped(paragraph) -> bool:
    style_name = (getattr(getattr(paragraph, 'style', None), 'name', '') or '').strip()
    if style_name.startswith('Heading') or style_name in ('Title', 'Subtitle'):
        return True
    text = (paragraph.text or '').strip()
    if not text:
        return False
    if _NUMBERED_HEADING_RE.match(text):
        return True
    for run in paragraph.runs:
        if not run.bold:
            continue
        size = run.font.size
        if size is not None and size.pt is not None and size.pt >= 14:
            return True
    return False


def _is_section_heading(paragraph) -> bool:
    """A section-level heading: Heading 2+ or a numbered list-style title.

    Excludes Heading 1 / Title (the document banner) so an AFTER_SAMPLES
    anchor doesn't latch onto a title like
    "Fiche de Réception d'Échantillons".
    """
    if not _is_heading_shaped(paragraph):
        return False
    style_name = (getattr(getattr(paragraph, 'style', None), 'name', '') or '').strip()
    if style_name in ('Heading 1', 'Title'):
        return False
    return True


def _matches_anchor(paragraph, keywords) -> bool:
    if not _is_section_heading(paragraph):
        return False
    text = (paragraph.text or '').lower()
    return any(kw.lower() in text for kw in keywords)


def _find_section_end(doc: DocumentType, anchor_heading) -> Optional[object]:
    """Return the paragraph that starts the NEXT section, or None at end.

    Used to compute "after the X section" positions: the block is inserted
    immediately before the next section-level heading that follows
    ``anchor_heading``. If no such next section exists the section runs to
    end-of-document and the caller falls back to appending.

    Compares via the underlying ``_element`` (XML node identity) rather
    than Python ``is`` because ``doc.paragraphs`` constructs fresh
    ``Paragraph`` wrappers on every access — two wrappers around the
    same XML element are never the same object.
    """
    anchor_elem = anchor_heading._element
    seen = False
    for p in doc.paragraphs:
        if not seen:
            if p._element is anchor_elem:
                seen = True
            continue
        if _is_section_heading(p):
            return p
    return None


def _find_anchor_for_position(doc: DocumentType, position: str):
    """Resolve a position label to an insertion anchor.

    Returns a tuple ``(reference_paragraph, where)`` where ``where`` is
    ``'before'``, ``'after'``, or ``'end'``. The caller uses ``where`` to
    decide whether to ``addprevious``, ``addnext`` or append a new paragraph.
    Falls back to ``(None, 'end')`` when a semantic anchor can't be found.
    """
    paragraphs = list(doc.paragraphs)
    if not paragraphs:
        return None, 'end'

    if position == 'TOP':
        # First heading-shaped paragraph (e.g. the H1 title). Insert AFTER
        # it so the block lands under the title but above the first section.
        for p in paragraphs:
            if _is_heading_shaped(p):
                return p, 'after'
        # Pathological — no headings at all. Drop the block at the start.
        return paragraphs[0], 'before'

    if position == 'BOTTOM':
        return None, 'end'

    if position == 'BEFORE_FOOTER':
        # The "Document généré automatiquement par PLAGENOR 4.0 …" line is
        # the very last non-empty paragraph in the body. Insert above it.
        for p in reversed(paragraphs):
            if (p.text or '').strip():
                return p, 'before'
        return None, 'end'

    if position in _ANCHOR_KEYWORDS:
        keywords = _ANCHOR_KEYWORDS[position]
        for p in paragraphs:
            if _matches_anchor(p, keywords):
                next_heading = _find_section_end(doc, p)
                if next_heading is not None:
                    return next_heading, 'before'
                return None, 'end'
        return None, 'end'

    return None, 'end'


def add_paragraph_after(reference_paragraph, text: str = '', *, bold: bool = False,
                        font_size_pt: Optional[float] = None) -> object:
    """Insert a new paragraph immediately after ``reference_paragraph``.

    python-docx doesn't expose a public API for mid-body insertion, so we
    manipulate the underlying OOXML element directly. Returns a
    :class:`docx.text.paragraph.Paragraph` wrapping the new element so the
    caller can add runs / set style.
    """
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph

    new_p = OxmlElement('w:p')
    reference_paragraph._element.addnext(new_p)
    paragraph = Paragraph(new_p, reference_paragraph._parent)
    if text:
        run = paragraph.add_run(text)
        if bold:
            run.bold = True
        if font_size_pt is not None:
            run.font.size = Pt(font_size_pt)
    return paragraph


def add_paragraph_before(reference_paragraph, text: str = '', *, bold: bool = False,
                         font_size_pt: Optional[float] = None) -> object:
    """Insert a new paragraph immediately before ``reference_paragraph``."""
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph

    new_p = OxmlElement('w:p')
    reference_paragraph._element.addprevious(new_p)
    paragraph = Paragraph(new_p, reference_paragraph._parent)
    if text:
        run = paragraph.add_run(text)
        if bold:
            run.bold = True
        if font_size_pt is not None:
            run.font.size = Pt(font_size_pt)
    return paragraph


# Legacy IBTIKAR forms — label-text substitution ---------------------------

# The egtp_*.docx templates ship as blank printable forms: each profile
# field is encoded as a French label followed by ``: *`` and (optionally) a
# Microsoft Word "instructional text" hint such as ``Nom complet du
# demandeur``. They contain no ``{{KEY}}`` markers, so the standard
# substitution pass leaves them empty. This map injects the actual values
# inline while preserving the original label text and asterisk.
#
# Pattern semantics:
# * The regex matches the FULL "Label : *  hint" line including any trailing
#   instructional text — Word inserts a non-breaking space (``\xa0``) and
#   sometimes a tab between ``*`` and the hint.
# * The replacement format keeps the label and asterisk so the form still
#   reads as an official fillable form (the value sits in the hint slot).
# * If the field is empty (e.g. a guest submission with no laboratory),
#   the original hint is preserved so the line still makes sense.

_LEGACY_LABEL_RULES = [
    # (regex, field_map_key, label_template)
    # 1. Personal info section
    (re.compile(r'Nom et prénom\s*:\s*\*[^\n]*', re.IGNORECASE),
     'FULL_NAME', 'Nom et prénom : * {value}'),
    (re.compile(r'(?:Université\s*/\s*École|Établissement)\s*:\s*\*[^\n]*', re.IGNORECASE),
     'ETABLISSEMENT', 'Université / École : * {value}'),
    (re.compile(r'Laboratoire\s*:\s*\*?[^\n]*', re.IGNORECASE),
     'LABORATORY', 'Laboratoire : {value}'),
    (re.compile(r'Fonction\s*/\s*Poste\s*:\s*\*[^\n]*', re.IGNORECASE),
     'STUDENT_LEVEL', 'Fonction / Poste : * {value}'),
    (re.compile(r'Adresse\s+e-?mail\s*:\s*\*[^\n]*', re.IGNORECASE),
     'EMAIL', 'Adresse e-mail : * {value}'),
    (re.compile(r'Numéro de téléphone(?:\s+du\s+demandeur)?\s*:\s*\*[^\n]*', re.IGNORECASE),
     'PHONE', 'Numéro de téléphone : * {value}'),
    # Optional DGRSDT identifier (only inserted if the requester filled it)
    (re.compile(r'(?:Identifiant\s+IBTIKAR|ID(?:GRSDT)?)\s*:\s*\*?[^\n]*', re.IGNORECASE),
     'IBTIKAR_ID', 'Identifiant IBTIKAR : {value}'),
    # 2. Request details
    (re.compile(r'Titre du projet\s*:\s*\*[^\n]*', re.IGNORECASE),
     'TITLE', 'Titre du projet : * {value}'),
    (re.compile(r'Directeur de recherche\s*:\s*\*[^\n]*', re.IGNORECASE),
     'SUPERVISOR', 'Directeur de recherche : * {value}'),
    # 3. Analysis framework (PFE / 1275 / 008 / Projet de doctorat / Autre)
    (re.compile(r"Cadre de l[’']analyse\s*:\s*\*?[^\n]*", re.IGNORECASE),
     'ANALYSIS_FRAME', 'Cadre de l’analyse : * {value}'),
]

# Plain string substitutions for the request-number header and date stamps
# that aren't in label form.
_LEGACY_TEXT_RULES = [
    # (find_pattern, replacement using field_map keys)
    (re.compile(r'……\.+/(\d{4})/IBTIKAR/PLAGENOR/ESSBO'),
     r'{DISPLAY_ID}/\1/IBTIKAR/PLAGENOR/ESSBO'),
]


def apply_legacy_label_substitution(doc: DocumentType, field_map: Mapping[str, str]) -> None:
    """Fill the literal-label fields in legacy egtp_*.docx IBTIKAR forms.

    Idempotent. Skips fields whose value is empty/N/A so we never overwrite
    a meaningful hint with an empty string (the requester or operator can
    still complete those by hand). Only applied when the active template
    is one of the egtp_*.docx forms — the generic and programmatic paths
    use ``{{KEY}}`` substitution instead and don't need this layer.
    """

    def _substitute(paragraph) -> None:
        joined = ''.join(r.text or '' for r in paragraph.runs)
        if not joined:
            return
        original = joined

        # Apply label-based rules (replace the whole label line).
        for pattern, key, template in _LEGACY_LABEL_RULES:
            value = field_map.get(key, '')
            if not value or value in ('N/A', 'Non défini', 'Non assigné'):
                continue
            joined = pattern.sub(template.format(value=value), joined)

        # Apply free-text rules (request number etc.).
        for pattern, template in _LEGACY_TEXT_RULES:
            try:
                joined = pattern.sub(template.format(**field_map), joined)
            except (KeyError, IndexError):
                pass

        if joined == original:
            return
        if paragraph.runs:
            paragraph.runs[0].text = joined
            for r in paragraph.runs[1:]:
                r.text = ''

    def _visit(paragraphs):
        for p in paragraphs:
            _substitute(p)

    def _visit_tables(tables):
        for t in tables:
            for row in t.rows:
                for cell in row.cells:
                    _visit(cell.paragraphs)
                    _visit_tables(cell.tables)

    _visit(doc.paragraphs)
    _visit_tables(doc.tables)
    for section in doc.sections:
        for header in (section.header, section.first_page_header, section.even_page_header):
            if header is not None:
                _visit(header.paragraphs)
                _visit_tables(header.tables)
        for footer in (section.footer, section.first_page_footer, section.even_page_footer):
            if footer is not None:
                _visit(footer.paragraphs)
                _visit_tables(footer.tables)


# Legacy IBTIKAR forms — sample-table population ---------------------------

def _norm_token(text: str) -> str:
    """Lowercase, strip accents and non-alphanumerics for fuzzy header match."""
    import unicodedata
    text = unicodedata.normalize('NFKD', text or '')
    text = ''.join(c for c in text if not unicodedata.combining(c))
    return re.sub(r'[^a-z0-9]', '', text.lower())


# Synonym groups mapping a submitted sample key (or its tokens) to the
# branded-form column it belongs in. Each entry: canonical -> set of tokens
# that may appear either in the submitted key or the DOCX column header.
_SAMPLE_COLUMN_SYNONYMS = [
    ('code', {'code', 'samplecode', 'echantillon', 'sample', 'refechantillon', 'reference', 'ref'}),
    ('origine', {'origine', 'origin', 'origineadn', 'origindna', 'source'}),
    ('typeadn', {'typeadn', 'typednad', 'dnatype', 'type', 'typedna', 'adn', 'dna'}),
    ('extraction', {'methode', 'extraction', 'methodeextraction', 'extractionmethod', 'method'}),
    ('gene', {'gene', 'genecible', 'targetgene', 'cible', 'target', 'amplicon', 'taille'}),
    ('amorces', {'amorces', 'primers', 'sequences', 'sequence', 'primer'}),
    ('tm', {'tm', 'temperature', 'meltingtemp', 'tmc'}),
    ('remarques', {'remarques', 'remarks', 'notes', 'note', 'commentaire', 'comment', 'observations'}),
    ('numero', {'n', 'no', 'num', 'numero', 'number', 'ndordre'}),
]


def _canonical_for(token: str) -> Optional[str]:
    """Return the canonical column group whose synonyms contain ``token``.

    Exact matches win first (so a one-letter header like ``N`` maps to
    ``numero``, not to ``code`` via a stray substring hit inside
    ``echantillon``). Substring matching is the fallback and only fires for
    tokens of 4+ chars against synonyms of 4+ chars, which is what compound
    headers like ``genecibletaille`` need.
    """
    if not token:
        return None
    for canonical, toks in _SAMPLE_COLUMN_SYNONYMS:
        if token in toks:
            return canonical
    if len(token) >= 4:
        for canonical, toks in _SAMPLE_COLUMN_SYNONYMS:
            for t in toks:
                if len(t) >= 4 and (t in token or token in t):
                    return canonical
    return None


def _set_cell_text(cell, value: str) -> None:
    """Write ``value`` into a table cell, preserving the cell's first-run style."""
    if not cell.paragraphs:
        cell.add_paragraph(value)
        return
    para = cell.paragraphs[0]
    if para.runs:
        para.runs[0].text = value
        for r in para.runs[1:]:
            r.text = ''
    else:
        para.add_run(value)


def populate_legacy_sample_table(doc: DocumentType, request_obj) -> None:
    """Fill the branded IBTIKAR form's sample table from the submitted data.

    The egtp_*.docx forms ship as blank printable grids. Best practice is for
    the generated document to carry everything the system already knows, so
    the lab receives a complete form rather than re-typing the sample list by
    hand. This maps each submitted ``sample_table`` row into the matching
    DOCX columns (by accent-insensitive synonym matching on the header) and
    fills the data rows, adding rows when there are more samples than the
    template pre-printed.

    Idempotent and defensive: any structural surprise (no recognisable table,
    no samples) leaves the document untouched.
    """
    samples = getattr(request_obj, 'sample_table', None)
    if not samples or not isinstance(samples, list):
        return

    # The YAML sample_table.columns carry, per service, the exact French
    # label that appears on the fiche. We map each DOCX header to the YAML
    # column whose label/name best matches it — so arbitrary service-specific
    # columns (storage conditions, culture medium…) fill too, not just the
    # known bio columns. Falls back to the synonym table when no YAML column
    # definition is available.
    service_code = getattr(getattr(request_obj, 'service', None), 'code', '') or ''
    yaml_columns = []  # list of (name, normalized_label, normalized_name)
    try:
        from core.registry import get_service_def
        sdef = get_service_def(service_code) or {}
        for col in (sdef.get('sample_table', {}) or {}).get('columns', []) or []:
            cname = col.get('name')
            if not cname:
                continue
            yaml_columns.append((
                cname,
                _label_norm(col.get('label', cname)),
                _label_norm(cname.replace('_', ' ')),
            ))
    except Exception:
        yaml_columns = []

    def header_to_field(header_text: str):
        """Return the submitted-sample key this DOCX header maps to."""
        h = _label_norm(header_text)
        if h in ('n', 'no', 'num', 'numero'):
            return '__numero__'
        if not h:
            return None
        # 1) best YAML column by label/name overlap
        best, best_score = None, 0
        hw = set(h.split())
        for cname, lnorm, nnorm in yaml_columns:
            for cand in (lnorm, nnorm):
                if not cand:
                    continue
                if cand == h or cand in h or h in cand:
                    score = 1000 + len(cand)
                else:
                    cw = set(cand.split())
                    score = len(hw & cw)
                    if score == 0:
                        continue
                if score > best_score:
                    best_score, best = score, cname
        if best is not None and best_score > 0:
            return best
        # 2) synonym fallback (canonical key); samples may be keyed that way
        canon = _canonical_for(_norm_token(header_text))
        return canon

    # 1) Find the sample table: the largest table with >=3 columns whose
    #    header maps to >=2 fields (skips the 2-col validation/visa table).
    target_table = None
    header_map = None  # {col_index: field_key}
    for table in sorted(doc.tables, key=lambda t: -len(t.columns)):
        if len(table.rows) < 2 or len(table.columns) < 3:
            continue
        hdr = {}
        for ci, cell in enumerate(table.rows[0].cells):
            f = header_to_field(cell.text)
            if f and (f == '__numero__' or f not in hdr.values()):
                hdr[ci] = f
        mapped = [v for v in hdr.values() if v != '__numero__']
        if len(mapped) >= 2:
            target_table, header_map = table, hdr
            break
    if target_table is None or not header_map:
        return

    # 2) Normalise each submitted sample to {field_key: value}. Keys are the
    #    YAML column names the online form posts; also accept synonym keys.
    def sample_to_fields(sample: dict) -> dict:
        out = {}
        if not isinstance(sample, dict):
            return out
        valid = {c[0] for c in yaml_columns}
        for key, val in sample.items():
            if val in (None, '', [], {}):
                continue
            sval = str(val)
            if key in valid:
                out[key] = sval
            else:
                canon = _canonical_for(_norm_token(str(key)))
                if canon and canon not in out:
                    out[canon] = sval
        return out

    rows_data = [sample_to_fields(s) for s in samples]
    rows_data = [r for r in rows_data if r]
    if not rows_data:
        return

    # 3) Fill data rows (after the header), adding rows when needed.
    data_rows = target_table.rows[1:]
    for idx, rdata in enumerate(rows_data):
        row = data_rows[idx] if idx < len(data_rows) else target_table.add_row()
        cells = row.cells
        for ci, field in header_map.items():
            if ci >= len(cells):
                continue
            if field == '__numero__':
                _set_cell_text(cells[ci], str(idx + 1).zfill(2))
                continue
            value = rdata.get(field)
            if value:
                _set_cell_text(cells[ci], value)


# Legacy IBTIKAR forms — generic question-filler ---------------------------

# Hints that mark an "empty" answer slot in a Word form paragraph. When we see
# any of these immediately after a label, we replace them with the requester's
# answer.
_EMPTY_ANSWER_HINTS = [
    "choisissez un élément",
    "choisissez un element",
    "cliquez ou appuyez ici",
    "click or tap here",
    "select an item",
]

# Regex finds: "<question label> : <hint><optional trailing words>".
# Group 1 = label, group 2 = the hint. Trailing text after the hint is
# allowed because Word writes "Cliquez ou appuyez ici pour entrer du texte."
# (the hint is not always at the very end of the line).
_QUESTION_LINE_RE = re.compile(
    r'^(.{8,160}?)\s*:\s*(' +
    '|'.join(re.escape(h) for h in _EMPTY_ANSWER_HINTS) +
    r')[^\n]*$',
    re.IGNORECASE | re.DOTALL,
)


def _label_norm(text: str) -> str:
    """Aggressive normalisation for fuzzy label matching."""
    import unicodedata
    text = unicodedata.normalize('NFKD', text or '')
    text = ''.join(c for c in text if not unicodedata.combining(c))
    # drop bracketed asides like "(chaque kit est associé à un tarif spécifique)"
    text = re.sub(r'\([^)]*\)', ' ', text)
    return re.sub(r'[^a-z0-9]+', ' ', text.lower()).strip()


def _build_param_label_index(service_params: dict, service_code: str) -> list:
    """Return [(normalized_label, value, original_label, original_name), ...]
    drawn from the YAML registry so we can map the DOCX questions to answers
    no matter how the labels are phrased.

    Multiple entries per param are emitted when the YAML provides bilingual
    labels (``label_fr`` / ``label_en``), so a French DOCX line can match an
    English-labelled param via its French alias and vice versa.
    """
    if not service_params or not service_code:
        return []
    try:
        from core.registry import get_service_def
        sdef = get_service_def(service_code) or {}
    except Exception:
        sdef = {}
    yaml_params = sdef.get('parameters') or []
    by_name = {p['name']: p for p in yaml_params}
    out = []
    seen_norm = set()
    for name, value in service_params.items():
        if value in (None, '', [], {}):
            continue
        if isinstance(value, bool):
            value = 'Oui' if value else 'Non'
        elif not isinstance(value, str):
            value = str(value)
        p = by_name.get(name) or {}
        labels = [
            p.get('label_fr'),
            p.get('label'),
            p.get('label_en'),
            name.replace('_', ' '),
        ]
        for raw in labels:
            if not raw:
                continue
            norm = _label_norm(raw)
            if not norm or norm in seen_norm:
                continue
            seen_norm.add(norm)
            out.append((norm, value, raw, name))
    return out


def _fuzzy_pick(question_norm: str, candidates: list, used: set):
    """Pick the best (label_norm, value, ...) candidate for ``question_norm``.

    Scoring: exact match wins; otherwise count word overlap. A candidate is
    skipped if it was already used for an earlier question. Ties resolve in
    insertion order. Returns the candidate tuple or ``None`` when no overlap.
    """
    q_words = set(question_norm.split())
    if not q_words:
        return None
    best = None
    best_score = 0
    for cand in candidates:
        cand_norm = cand[0]
        if cand_norm in used:
            continue
        c_words = set(cand_norm.split())
        if not c_words:
            continue
        if cand_norm == question_norm or cand_norm in question_norm or question_norm in cand_norm:
            score = len(c_words) + 1000  # exact/containment beats overlap
        else:
            overlap = len(q_words & c_words)
            # require at least one substantive word in common (>= 3 chars)
            if not any(w in q_words for w in c_words if len(w) >= 3 and w in q_words):
                continue
            score = overlap
        if score > best_score:
            best_score = score
            best = cand
    return best


def populate_legacy_param_questions(doc: DocumentType, request_obj) -> None:
    """Fill every "Choisissez un élément" / "Cliquez…" slot in the legacy
    IBTIKAR forms with the answer the requester provided online.

    The match is fuzzy on labels: the DOCX question (e.g. "Type de kit PCR
    que vous souhaitez utiliser") is mapped to the YAML param whose label
    overlaps the most ("PCR kit"). Each param is used at most once, so two
    similar questions don't both grab the same answer.

    Idempotent and defensive: paragraphs with no recognisable hint are
    skipped; tables are walked recursively; styling is preserved by
    rewriting only the first run.
    """
    service_code = getattr(getattr(request_obj, 'service', None), 'code', '') or ''
    candidates = _build_param_label_index(
        getattr(request_obj, 'service_params', None) or {},
        service_code,
    )
    if not candidates:
        return
    used: set = set()

    def _fill_paragraph(paragraph) -> None:
        joined = ''.join(r.text or '' for r in paragraph.runs)
        if not joined:
            return
        m = _QUESTION_LINE_RE.match(joined.strip())
        if not m:
            return
        label_text, _hint = m.group(1).strip(), m.group(2)
        # Skip the personal-info labels — those go through the label-rules path
        # and we don't want a double substitution.
        if re.search(r'(nom et prénom|adresse e-?mail|téléphone|laboratoire|fonction|directeur|titre du projet|cadre de l[’\']analyse)',
                     label_text, re.IGNORECASE):
            return
        picked = _fuzzy_pick(_label_norm(label_text), candidates, used)
        if not picked:
            return
        norm, value, _orig_label, _name = picked
        used.add(norm)
        new_text = f"{label_text} : {value}"
        if paragraph.runs:
            paragraph.runs[0].text = new_text
            for r in paragraph.runs[1:]:
                r.text = ''

    def _visit_paragraphs(paragraphs):
        for p in paragraphs:
            _fill_paragraph(p)

    def _visit_tables(tables):
        for t in tables:
            for row in t.rows:
                for cell in row.cells:
                    _visit_paragraphs(cell.paragraphs)
                    _visit_tables(cell.tables)

    _visit_paragraphs(doc.paragraphs)
    _visit_tables(doc.tables)
