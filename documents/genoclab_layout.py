"""GENOCLAB commercial documents layout (quote + invoice).

Replicates the SAIDAL-style "Facture Proforma" model the platform uses
for commercial billing: GENOCELAB logo top-left, two-column header
(issuer info on the left, client info on the right), prestation table
with columns Prestation / Quantité / Prix unitaire DA / Montant DA,
HT / VAT / TTC totals, and a legal footer with the amount-in-words
line plus the registered-office block.

Every text-only element (issuer name, NIF, bank accounts, footer
legal text, …) is read from PlatformContent so the SuperAdmin can
edit it via /dashboard/home/content/update/ without touching code.
The defaults match the model file supplied by the owner.
"""
from __future__ import annotations

from decimal import Decimal
from pathlib import Path
from typing import Iterable, Optional

from docx.document import Document as DocumentType
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from documents.docx_helpers import (
    BRAND_DARK,
    BRAND_FONT,
    BRAND_MUTED,
    SIZE_BODY,
    SIZE_CAPTION,
    SIZE_H1,
    SIZE_H2,
    _GENOCLAB_LOGO,
)


# Logo colours — sampled from the GENOCLAB asset itself (see comment at
# the bottom of this module for the sampling script). Used in place of
# the PLAGENOR indigo for everything on commercial documents so the
# devis/facture visually descend from the GENOCLAB brand instead of the
# academic one.
GCL_NAVY = RGBColor(0x18, 0x30, 0x60)   # the "CLAB" + DNA helix dark blue
GCL_TEAL = RGBColor(0x18, 0xA8, 0xA8)   # the "GENO" + building turquoise
GCL_TEAL_TINT = 'D8F0F0'                # very-light teal fill for header row


# Default values for every editable string. Kept in lock-step with the
# SAIDAL Proforma model supplied by the platform owner; the SuperAdmin
# can override any of them from the CMS without code changes.
CMS_DEFAULTS = {
    'genoclab_issuer_name':       "École Supérieure en Sciences Biologiques d'Oran (ESSBO)",
    'genoclab_issuer_address1':   "BP 1042 SAIM MOHAMED,",
    'genoclab_issuer_address2':   "Cité Emir Abdelkader (EX-INESSMO)",
    'genoclab_issuer_address3':   "31000 Oran",
    'genoclab_issuer_treasury':   "Cpte Trésor : 00831001131000208471",
    'genoclab_issuer_nif':        "N.I.F : 415020000310784",
    'genoclab_issuer_ccp':        "Cpte CCP Agent comptable de l'ESSBO : 007999990000",
    'genoclab_issuer_phone':      "Téléphone / Fax : +213 41 24 63 59",
    'genoclab_quote_title':       "Facture Proforma",
    'genoclab_invoice_title':     "Facture",
    'genoclab_footer_legal':      (
        "Arrêtée la présente facture à la somme de "
        "____________________________________________________________ Dinars Algériens."
    ),
    'genoclab_footer_office':     (
        "Siège social — BP 1042 SAIM MOHAMED, Cité Emir Abdelkader (EX-INESSMO), 31000 Oran"
    ),
    'genoclab_footer_contact':    (
        "École Supérieure en Sciences Biologiques d'Oran (ESSBO) · "
        "https://essb-oran.edu.dz/"
    ),
    'genoclab_vat_rate':          "0.19",
}


def cms_get(key: str, default: str = '') -> str:
    """Read an editable string from PlatformContent. Falls back to the
    CMS_DEFAULTS dict, then to the provided default."""
    try:
        from core.models import PlatformContent
        obj = PlatformContent.objects.filter(key=key, lang='fr').first()
        if obj and (obj.value or '').strip():
            return obj.value
    except Exception:
        pass
    return CMS_DEFAULTS.get(key, default)


# ── Amount-to-words ────────────────────────────────────────────────────────
# Tiny standalone French converter. We avoid a num2words dependency to
# keep requirements.txt unchanged. Algerian French uses the Belgian /
# Swiss style "septante / nonante" for some, but the official invoicing
# convention sticks to the French standard ("soixante-dix / quatre-vingt-
# dix"), so we follow that.

_UNITS = ('zéro', 'un', 'deux', 'trois', 'quatre', 'cinq', 'six', 'sept',
          'huit', 'neuf', 'dix', 'onze', 'douze', 'treize', 'quatorze',
          'quinze', 'seize', 'dix-sept', 'dix-huit', 'dix-neuf')
_TENS = ('', '', 'vingt', 'trente', 'quarante', 'cinquante', 'soixante',
         'soixante', 'quatre-vingt', 'quatre-vingt')


def _two_digits(n: int) -> str:
    if n < 20:
        return _UNITS[n]
    tens, units = divmod(n, 10)
    base = _TENS[tens]
    if tens in (7, 9):
        # 70-79 → soixante-dix, soixante-onze, …
        # 90-99 → quatre-vingt-dix, quatre-vingt-onze, …
        return f"{base}-{_UNITS[10 + units]}" if units else f"{base}-{_UNITS[10]}"
    if tens == 8 and units == 0:
        return 'quatre-vingts'
    if units == 1 and tens in (2, 3, 4, 5, 6):
        return f"{base} et un"
    if units == 0:
        return base
    return f"{base}-{_UNITS[units]}"


def _three_digits(n: int, *, followed_by_multiplier: bool = False) -> str:
    """French 0-999 written out.

    ``followed_by_multiplier`` should be True when this triplet is the
    multiplier of a higher unit (mille / million / milliard). It strips
    the plural 's' on 'cent' since "cent" loses its plural mark whenever
    it precedes another numeral name (deux cent mille — not "deux cents
    mille").
    """
    if n == 0:
        return ''
    if n < 100:
        return _two_digits(n)
    hundreds, rest = divmod(n, 100)
    if hundreds == 1:
        head = 'cent'
    else:
        head = f"{_UNITS[hundreds]} cent"
        # Plural "s" on "cent" only when it isn't followed by another
        # number (so "deux cents" but "deux cent mille").
        if rest == 0 and not followed_by_multiplier:
            head += 's'
    if rest:
        return f"{head} {_two_digits(rest)}"
    return head


def amount_in_words_fr(amount) -> str:
    """Return a French textual representation of ``amount`` in DZD.

    Example:
        amount_in_words_fr(11900) -> "onze mille neuf cents"
        amount_in_words_fr(123)   -> "cent vingt-trois"
        amount_in_words_fr(1000.50) -> "mille et 50/100"

    Caller is responsible for appending the currency name. Cents (chimes)
    are rendered as a fraction so the legal phrasing stays unambiguous.
    """
    try:
        amt = float(amount)
    except (TypeError, ValueError):
        return ''
    if amt < 0:
        return f"moins {amount_in_words_fr(-amt)}"
    integer_part = int(amt)
    cents = round((amt - integer_part) * 100)

    if integer_part == 0:
        words = 'zéro'
    else:
        # Split into milliards / millions / mille / units.
        groups = []
        for power, label in [(10**9, 'milliard'), (10**6, 'million'),
                             (10**3, 'mille'), (1, '')]:
            count = integer_part // power
            integer_part = integer_part % power
            if count == 0:
                continue
            if label == 'mille':
                # "mille" is invariable — never an 's'. And "un mille" is
                # wrong; just write "mille" when count == 1.
                if count == 1:
                    groups.append('mille')
                else:
                    groups.append(f"{_three_digits(count, followed_by_multiplier=True)} mille")
            elif label == '':
                groups.append(_three_digits(count))
            else:
                # million / milliard — keep the 's' when plural.
                plural = 's' if count > 1 else ''
                groups.append(f"{_three_digits(count, followed_by_multiplier=True)} {label}{plural}")
        words = ' '.join(g for g in groups if g)

    if cents:
        words += f" et {cents:02d}/100"
    return words


# ── End amount-to-words helpers ────────────────────────────────────────────


def _money(value, currency: str = 'DA') -> str:
    """Format a numeric amount with French thousands grouping."""
    try:
        x = float(value or 0)
    except (TypeError, ValueError):
        return str(value)
    return f"{x:,.2f}".replace(',', ' ').replace('.', ',') + f" {currency}"


def _money_int(value) -> str:
    """Whole-DA formatting for the unit-price column where decimals are noise."""
    try:
        x = float(value or 0)
    except (TypeError, ValueError):
        return str(value)
    if x == int(x):
        return f"{int(x):,}".replace(',', ' ')
    return f"{x:,.2f}".replace(',', ' ').replace('.', ',')


def _set_cell_text(cell, text: str, *, bold: bool = False, size: int = SIZE_BODY,
                   color=None, align=None) -> None:
    """Replace a table cell's content with one styled run."""
    cell.text = ''  # wipe whatever was there
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(text or '')
    run.font.name = BRAND_FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _shade_cell(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)


def add_genoclab_header(doc: DocumentType, *, title: str, doc_number: str,
                        doc_date: str, client_name: str = '',
                        client_lines: Optional[Iterable[str]] = None) -> None:
    """Write the GENOCLAB document header: logo top-left, big title,
    then a two-column block — issuer (CMS-editable) on the left, client
    coordinates on the right, with the date / document number below.
    """
    # Logo. Sized at ~6 cm wide, leaves comfortable whitespace next to it.
    if _GENOCLAB_LOGO.exists():
        p = doc.add_paragraph()
        run = p.add_run()
        run.add_picture(str(_GENOCLAB_LOGO), width=Cm(7))

    # Document title (e.g. "Facture Proforma" or "Facture") — large, brand colour.
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(title)
    run.font.name = BRAND_FONT
    run.font.size = Pt(SIZE_H1 + 4)  # 20 pt
    run.font.bold = True
    run.font.color.rgb = GCL_NAVY

    # Two-column header: issuer left, client right.
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    # No borders, no fill — just a layout grid.
    _clear_table_borders(table)

    issuer_cell = table.rows[0].cells[0]
    issuer_cell.width = Cm(10)
    client_cell = table.rows[0].cells[1]
    client_cell.width = Cm(7)

    # Issuer block (left)
    _multi_para(issuer_cell, [
        (cms_get('genoclab_issuer_name'),     {'bold': True, 'size': SIZE_BODY + 1}),
        (cms_get('genoclab_issuer_address1'), {}),
        (cms_get('genoclab_issuer_address2'), {}),
        (cms_get('genoclab_issuer_address3'), {}),
        (cms_get('genoclab_issuer_treasury'), {'size': SIZE_CAPTION + 1, 'color': BRAND_MUTED}),
        (cms_get('genoclab_issuer_nif'),      {'size': SIZE_CAPTION + 1, 'color': BRAND_MUTED}),
        (cms_get('genoclab_issuer_ccp'),      {'size': SIZE_CAPTION + 1, 'color': BRAND_MUTED}),
        (cms_get('genoclab_issuer_phone'),    {'size': SIZE_CAPTION + 1, 'color': BRAND_MUTED}),
    ])

    # Client block (right)
    client_paras = [("Client", {'bold': True, 'color': GCL_NAVY})]
    if client_name:
        client_paras.append((client_name, {'bold': True, 'size': SIZE_BODY + 1}))
    for line in (client_lines or []):
        if line:
            client_paras.append((line, {}))
    client_paras.append(("", {}))  # spacer
    client_paras.append((f"Date : {doc_date}", {'bold': True}))
    client_paras.append((f"N° : {doc_number}", {'bold': True}))
    _multi_para(client_cell, client_paras)


def _multi_para(cell, items) -> None:
    """Stuff multiple styled paragraphs into a single cell — first call
    overwrites, the rest are appended."""
    cell.text = ''
    p = cell.paragraphs[0]
    for i, (text, opts) in enumerate(items):
        if i > 0:
            p = cell.add_paragraph()
        run = p.add_run(text)
        run.font.name = BRAND_FONT
        run.font.size = Pt(opts.get('size', SIZE_BODY))
        run.font.bold = opts.get('bold', False)
        run.font.color.rgb = opts.get('color', BRAND_DARK)


def _clear_table_borders(table) -> None:
    """Remove every border from a python-docx table (used for layout
    tables that shouldn't look like data tables)."""
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = tcPr.find(qn('w:tcBorders'))
            if tcBorders is None:
                tcBorders = OxmlElement('w:tcBorders')
                tcPr.append(tcBorders)
            for side in ('top', 'left', 'bottom', 'right',
                         'insideH', 'insideV'):
                el = tcBorders.find(qn(f'w:{side}'))
                if el is None:
                    el = OxmlElement(f'w:{side}')
                    tcBorders.append(el)
                el.set(qn('w:val'), 'nil')


def add_prestation_table(doc: DocumentType, line_items,
                         *, vat_rate=None) -> None:
    """Render the SAIDAL-style prestation grid + subtotal / VAT / total
    block in the document. ``line_items`` is a list of dicts with keys
    ``label`` (or ``description``), ``quantity``, ``unit_price``,
    ``total``. Missing keys default to 0/empty. ``vat_rate`` is a float
    (e.g. 0.19); when None, read from CMS.
    """
    if vat_rate is None:
        try:
            vat_rate = float(cms_get('genoclab_vat_rate'))
        except (TypeError, ValueError):
            vat_rate = 0.19

    # Header row
    table = doc.add_table(rows=1 + len(line_items) + 3, cols=4)
    table.autofit = False
    headers = ['Prestation', 'Quantité', 'Prix unitaire (DA)', 'Montant (DA)']
    for i, h in enumerate(headers):
        _set_cell_text(table.rows[0].cells[i], h,
                       bold=True, color=GCL_NAVY,
                       align=(WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT))
        _shade_cell(table.rows[0].cells[i], GCL_TEAL_TINT)

    # Item rows
    subtotal_ht = Decimal('0')
    for r, item in enumerate(line_items, start=1):
        label = item.get('label') or item.get('description') or ''
        qty = item.get('quantity', 0) or 0
        unit_price = item.get('unit_price', 0) or 0
        total = item.get('total')
        if total is None:
            try:
                total = float(qty) * float(unit_price)
            except (TypeError, ValueError):
                total = 0
        try:
            subtotal_ht += Decimal(str(total))
        except Exception:
            pass
        cells = table.rows[r].cells
        _set_cell_text(cells[0], str(label))
        _set_cell_text(cells[1], str(qty), align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_text(cells[2], _money_int(unit_price), align=WD_ALIGN_PARAGRAPH.RIGHT)
        _set_cell_text(cells[3], _money_int(total), align=WD_ALIGN_PARAGRAPH.RIGHT)

    # Totals (subtotal HT, VAT, total TTC) — right-aligned, label in col 0-2 merged.
    vat_amount = float(subtotal_ht) * vat_rate
    total_ttc = float(subtotal_ht) + vat_amount

    def _total_row(row_idx, label, value, *, big=False):
        # Merge first three cells under the label so the layout reads
        # like the SAIDAL model.
        merged = table.rows[row_idx].cells[0]
        for ci in (1, 2):
            merged = merged.merge(table.rows[row_idx].cells[ci])
        _set_cell_text(
            merged, label,
            bold=True, size=SIZE_BODY + (1 if big else 0),
            color=GCL_NAVY if big else BRAND_DARK,
            align=WD_ALIGN_PARAGRAPH.RIGHT,
        )
        amount_cell = table.rows[row_idx].cells[3]
        _set_cell_text(
            amount_cell, _money_int(value),
            bold=True, size=SIZE_BODY + (1 if big else 0),
            color=GCL_NAVY if big else BRAND_DARK,
            align=WD_ALIGN_PARAGRAPH.RIGHT,
        )
        if big:
            _shade_cell(amount_cell, GCL_TEAL_TINT)

    base_row = 1 + len(line_items)
    _total_row(base_row,     "Sous-total HT",                                                       float(subtotal_ht))
    _total_row(base_row + 1, f"TVA ({int(round(vat_rate * 100))} %)",                                 vat_amount)
    _total_row(base_row + 2, "Total TTC",                                                              total_ttc, big=True)

    # Thin slate-200 borders on every cell of the table (data area).
    _apply_thin_borders(table)
    return total_ttc


def _apply_thin_borders(table) -> None:
    """Apply thin 1/2pt slate-200 borders on every cell of a python-docx
    table — used by add_prestation_table to demarcate the data grid."""
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = tcPr.find(qn('w:tcBorders'))
            if tcBorders is None:
                tcBorders = OxmlElement('w:tcBorders')
                tcPr.append(tcBorders)
            for side in ('top', 'left', 'bottom', 'right'):
                el = tcBorders.find(qn(f'w:{side}'))
                if el is None:
                    el = OxmlElement(f'w:{side}')
                    tcBorders.append(el)
                el.set(qn('w:val'), 'single')
                el.set(qn('w:sz'), '4')
                el.set(qn('w:color'), 'E2E8F0')


def add_genoclab_footer(doc: DocumentType, *, total_amount=None) -> None:
    """Legal & office block at the bottom of the document.

    Renders three CMS-editable blocks:
      1. The "Arrêtée la présente facture à la somme de …" legal phrase.
         When ``total_amount`` is supplied, the variable part (the
         underscored slot in the CMS default, or any "{amount_words}" /
         "{amount}" placeholder if the SuperAdmin uses one) is filled
         automatically with both the words and the figures, so the
         document is legally complete out of the box. The SuperAdmin
         keeps full control over the surrounding wording.
      2. The "Siège social …" line (CMS).
      3. The contact / website line (CMS).
    """
    doc.add_paragraph()  # spacer

    legal_text = cms_get('genoclab_footer_legal')
    if total_amount is not None:
        words = amount_in_words_fr(total_amount).strip()
        figures = _money_int(total_amount)
        words_upper = words[:1].upper() + words[1:] if words else ''

        if '{amount_words}' in legal_text or '{amount}' in legal_text:
            # Explicit placeholders — fill them in.
            legal_text = (
                legal_text
                .replace('{amount_words}', words_upper)
                .replace('{amount}', f"{figures} DA")
            )
        else:
            # Default phrasing has a long "____" placeholder we replace
            # with the words; we then append "(soit XXX DA)" so the
            # figure is unambiguous.
            import re as _re
            underscores = _re.search(r'_{5,}', legal_text)
            if underscores:
                legal_text = legal_text.replace(
                    underscores.group(0), words_upper)
            else:
                legal_text = legal_text.rstrip(' .') + f" {words_upper}"
            legal_text = legal_text.rstrip(' .') + f" (soit {figures} DA)."

    legal = doc.add_paragraph(legal_text)
    for run in legal.runs:
        run.font.name = BRAND_FONT
        run.font.size = Pt(SIZE_BODY)
        run.italic = True
        run.font.color.rgb = BRAND_DARK
    doc.add_paragraph()  # spacer

    office = doc.add_paragraph()
    run = office.add_run(cms_get('genoclab_footer_office'))
    run.font.name = BRAND_FONT
    run.font.size = Pt(SIZE_CAPTION + 1)
    run.font.color.rgb = BRAND_MUTED

    contact = doc.add_paragraph()
    run = contact.add_run(cms_get('genoclab_footer_contact'))
    run.font.name = BRAND_FONT
    run.font.size = Pt(SIZE_CAPTION + 1)
    run.font.color.rgb = BRAND_MUTED
