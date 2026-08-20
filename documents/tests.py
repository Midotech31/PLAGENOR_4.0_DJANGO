"""Tests for the DOCX generators.

They produce real .docx files (python-docx, no LibreOffice needed). We
generate a document from known data, reopen it, and assert the key values
made it into the file. Files land under MEDIA_ROOT/documents/ and are cleaned
up afterwards.
"""
import os
import shutil
import tempfile
from decimal import Decimal
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import MagicMock, patch

from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import SimpleTestCase, TestCase, override_settings
from docx import Document as _OpenDocx

_TEST_STORAGES = {
    'default': {'BACKEND': 'django.core.files.storage.FileSystemStorage'},
    'staticfiles': {'BACKEND': 'django.contrib.staticfiles.storage.StaticFilesStorage'},
}


def _docx_text(path):
    """Flatten all paragraph + table text of a .docx into one string."""
    doc = _OpenDocx(path)
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return '\n'.join(parts)


class InvoiceGeneratorTests(TestCase):
    def setUp(self):
        from accounts.models import User
        from core.models import Invoice
        self.client_user = User.objects.create(
            username='gen-client', role='CLIENT',
            first_name='Ada', last_name='Lovelace')
        self.invoice = Invoice.objects.create(
            invoice_number='GENOCLAB-INV-TEST-001',
            client=self.client_user,
            line_items=[{'description': 'Séquençage', 'quantity': 2,
                         'unit_price': 5000, 'total': 10000}],
            subtotal_ht=Decimal('10000'), vat_rate=Decimal('0.19'),
            vat_amount=Decimal('1900'), total_ttc=Decimal('11900'))
        self._paths = []

    def tearDown(self):
        for p in self._paths:
            if p and os.path.exists(p):
                os.remove(p)

    def test_invoice_docx_is_valid_and_contains_key_values(self):
        from documents.generators import generate_invoice_document
        path = generate_invoice_document(self.invoice)
        self._paths.append(path)
        self.assertTrue(path.endswith('.docx'))
        self.assertTrue(os.path.exists(path))
        text = _docx_text(path)  # opening also proves it's a valid docx
        self.assertIn('GENOCLAB-INV-TEST-001', text)
        self.assertIn('Séquençage', text)


class IbtikarFormGeneratorTests(TestCase):
    def setUp(self):
        from accounts.models import User
        from core.models import Request, Service
        self.service = Service.objects.create(
            code='GEN_IBK', name='Analyse génomique',
            channel_availability='IBTIKAR')
        self.requester = User.objects.create(
            username='gen-req', role='REQUESTER',
            first_name='Alan', last_name='Turing')
        self.req = Request.objects.create(
            channel='IBTIKAR', status='ASSIGNED', service=self.service,
            requester=self.requester, title='Demande test')
        self._paths = []

    def tearDown(self):
        for p in self._paths:
            if p and os.path.exists(p):
                os.remove(p)

    def test_ibtikar_form_is_valid_docx(self):
        from documents.generators import generate_ibtikar_form
        path = generate_ibtikar_form(self.req)
        self._paths.append(path)
        self.assertTrue(os.path.exists(path))
        text = _docx_text(path)  # must open without raising
        self.assertIn(self.req.display_id, text)

    def test_build_field_map_returns_dict(self):
        from documents.generators import build_field_map
        fm = build_field_map(self.req)
        self.assertIsInstance(fm, dict)
        self.assertTrue(len(fm) > 0)


class CompleteDocumentWorkflowTests(TestCase):
    """Exercise every production document family with realistic data."""

    def setUp(self):
        from accounts.models import MemberProfile, User
        from core.models import Request, Service

        self.media_root = tempfile.mkdtemp(prefix='plagenor-doc-tests-')
        self.settings_override = override_settings(MEDIA_ROOT=self.media_root)
        self.settings_override.enable()
        self.requester = User.objects.create_user(
            username='full-doc-requester', password='x', role='REQUESTER',
            first_name='Marie', last_name='Curie', email='marie@example.test',
            phone='0550000000', organization='ESSBO', laboratory='PLAGENOR',
            supervisor='Pr. Test', student_level='Doctorat', ibtikar_id='IBK-42',
            preferred_language='fr',
        )
        analyst_user = User.objects.create_user(
            username='full-doc-analyst', password='x', role='MEMBER',
            first_name='Ada', last_name='Lovelace', email='ada@example.test',
        )
        analyst = analyst_user.member_profile
        service = Service.objects.create(
            code='DOC-FULL', name='Séquençage complet',
            description='Analyse génomique approfondie', service_type='Sequencing',
            channel_availability='BOTH', turnaround_days=12,
            ibtikar_price=Decimal('2500'), genoclab_price=Decimal('4000'),
        )
        self.req = Request.objects.create(
            display_id='DOC-FULL-001', channel='GENOCLAB', status='QUOTE_DRAFT',
            service=service, requester=self.requester, assigned_to=analyst,
            title='Projet documentaire complet', description='Description détaillée',
            urgency='Urgent', budget_amount=Decimal('12000'),
            declared_ibtikar_balance=Decimal('20000'), quote_amount=Decimal('10000'),
            admin_validated_price=Decimal('9000'), ibtikar_external_code='DGRSDT-9',
            service_params={
                'analysis_frame': 'Projet de doctorat', 'quality_level': 'Premium',
                'paired_end': True, 'replicates': 3, 'empty': '',
            },
            sample_table=[
                {'sample_id': 'S1', 'organism_type': 'Bactérie', 'volume': 20},
                {'sample_id': 'S2', 'organism_type': 'Levure', 'volume': 30},
            ],
            quote_detail={
                'items': [
                    {'description': 'Extraction', 'quantity': 2, 'unit_price': 1500, 'total': 3000},
                    {'description': 'Séquençage', 'quantity': 1, 'unit_price': 7000, 'total': 7000},
                ],
                'admin_fees': 500,
                'report_fees': 250,
            },
        )
        self.paths = []

    def tearDown(self):
        self.settings_override.disable()
        shutil.rmtree(self.media_root, ignore_errors=True)

    def _assert_doc(self, generator, *needles):
        path = generator(self.req)
        self.paths.append(path)
        self.assertTrue(Path(path).exists())
        text = _docx_text(path)
        for needle in needles:
            self.assertIn(needle, text)
        return path

    def test_platform_note_quote_and_reception_documents(self):
        from documents.generators import (
            generate_platform_note, generate_quote, generate_reception_form,
        )

        self._assert_doc(generate_platform_note, 'DOC-FULL-001')
        self._assert_doc(generate_quote, 'GENOCLAB-DEV', 'Marie Curie')
        self._assert_doc(generate_reception_form, 'DOC-FULL-001')

    def test_programmatic_ibtikar_document_exercises_rich_fields(self):
        from documents.generators import generate_ibtikar_form

        with patch('documents.generators.Path.exists', return_value=False):
            path = generate_ibtikar_form(self.req)
        self.assertTrue(Path(path).exists())
        text = _docx_text(path)
        self.assertIn('Projet documentaire complet', text)
        self.assertIn('S1', text)
        self.assertIn('Premium', text)

    def test_stats_docx_and_excel_are_valid(self):
        from documents.generators import generate_stats_report
        from documents.stats_excel import generate_bilan_excel

        kpis = {
            'total': 12, 'completed': 7, 'in_progress': 3, 'rejected': 2,
            'completion_rate': 58.3, 'ibtikar_count': 8, 'genoclab_count': 4,
            'ibtikar_virtual_revenue': 100000, 'genoclab_revenue': 50000,
        }
        rows = [{'label': 'Séquençage', 'count': 7}, {'label': 'PCR', 'count': 5}]
        bundle = {
            'kpis': kpis, 'by_service': rows, 'by_status': rows,
            'by_wilaya': rows, 'by_organization': rows,
            'by_analysis_frame': rows, 'by_gender': rows,
            'trend': [{'month': '2026-08', 'count': 12}],
        }
        docx_path = generate_stats_report(
            bundle, {'date_from': '2026-01-01', 'date_to': '2026-08-31'},
            self.requester,
        )
        self.assertIn('Indicateurs principaux', _docx_text(docx_path))

        sections = [
            {'title': "Services / test", 'columns': ['Service', 'Demandes'],
             'rows': [['Séquençage', 7], ['PCR', 5]],
             'total_row': ['Total', 12]},
            {'title': "Services / test", 'columns': ['Service', 'Demandes'],
             'rows': [['Séquençage', 7], ['PCR', 5]],
             'total_row': ['Total', 12]},
        ]
        excel_path = generate_bilan_excel(
            {'granularity': 'month', 'kpis': kpis, 'sections': sections},
            {'date_from': '2026-01-01'}, self.requester,
        )
        self.assertTrue(Path(excel_path).exists())
        from openpyxl import load_workbook
        workbook = load_workbook(excel_path)
        self.assertIn('Synthèse', workbook.sheetnames)
        self.assertEqual(len(workbook.sheetnames), 3)

    def test_field_map_supports_guest_and_missing_service(self):
        from core.models import Request
        from documents.generators import build_field_map

        guest = Request.objects.create(
            display_id='GUEST-DOC-1', title='Guest', channel='GENOCLAB',
            submitted_as_guest=True, guest_name='Guest User',
            guest_email='guest@example.test', guest_phone='123',
            service_params={'flag': False, 'choices': ['A', 'B']},
        )
        field_map = build_field_map(guest)
        self.assertEqual(field_map['FULL_NAME'], 'Guest User')
        self.assertEqual(field_map['SERVICE_CODE'], 'N/A')
        self.assertIn('Flag', field_map['SERVICE_PARAMS'])


class DocxHelperUnitTests(SimpleTestCase):
    def test_placeholder_cleanup_styling_and_relative_insertion(self):
        from documents.docx_helpers import (
            add_brand_footer, add_paragraph_after, add_paragraph_before,
            apply_house_style, replace_placeholders, strip_unresolved_placeholders,
            style_brand_table,
        )

        doc = _OpenDocx()
        paragraph = doc.add_paragraph()
        paragraph.add_run('Hello {{NA')
        paragraph.add_run('ME}} / {{UNKNOWN}}')
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = '{{NAME}}'
        table.cell(0, 1).text = '{{UNKNOWN}}'
        table.cell(1, 0).text = 'Value'
        replace_placeholders(doc, {'NAME': 'PLAGENOR'})
        strip_unresolved_placeholders(doc)
        apply_house_style(doc)
        style_brand_table(table, accent='header')
        style_brand_table(table, accent='subtle')
        before = add_paragraph_before(paragraph, 'Before', bold=True, font_size_pt=11)
        after = add_paragraph_after(paragraph, 'After', bold=True, font_size_pt=11)
        add_brand_footer(doc, organisation='ESSBO')
        self.assertIn('PLAGENOR', table.cell(0, 0).text)
        self.assertNotIn('{{', '\n'.join(p.text for p in doc.paragraphs))
        self.assertEqual(before.text, 'Before')
        self.assertEqual(after.text, 'After')

    def test_legacy_matching_helpers_and_table_population(self):
        from documents.docx_helpers import (
            _canonical_for, _fuzzy_pick, _norm_token, populate_legacy_sample_table,
        )

        self.assertEqual(_norm_token('  Échantillon-N°  '), 'echantillonn')
        self.assertEqual(_canonical_for('samplecode'), 'code')
        candidates = [('quality level', 'quality_level', 'Premium')]
        self.assertEqual(_fuzzy_pick('quality level', candidates, set())[1], 'quality_level')

        doc = _OpenDocx()
        table = doc.add_table(rows=3, cols=3)
        for index, label in enumerate(('N°', 'Code échantillon', 'Origine')):
            table.cell(0, index).text = label
        request_obj = SimpleNamespace(
            sample_table=[{'sample_code': 'S-01', 'sample_origin': 'Sol'}],
            service=None,
        )
        populate_legacy_sample_table(doc, request_obj)
        flattened = '\n'.join(cell.text for row in table.rows for cell in row.cells)
        self.assertIn('S-01', flattened)


class PdfConverterTests(SimpleTestCase):
    def setUp(self):
        self.tmp = tempfile.mkdtemp(prefix='plagenor-pdf-tests-')
        self.docx = Path(self.tmp) / 'source.docx'
        self.docx.write_bytes(b'docx')

    def tearDown(self):
        shutil.rmtree(self.tmp, ignore_errors=True)

    @override_settings(DOCUMENT_PDF_ENABLED=False)
    def test_disabled_and_missing_inputs(self):
        from documents.pdf_converter import convert_docx_to_pdf

        self.assertEqual(convert_docx_to_pdf(self.docx), self.docx)
        with self.assertRaises(FileNotFoundError):
            convert_docx_to_pdf(Path(self.tmp) / 'missing.docx')

    @override_settings(DOCUMENT_PDF_ENABLED=True, DOCUMENT_PDF_BACKEND='spawn')
    def test_spawn_success_failure_timeout_and_missing_binary(self):
        from documents.pdf_converter import _convert_via_spawn, convert_docx_to_pdf

        out = Path(self.tmp) / 'out'
        out.mkdir()
        with patch('documents.pdf_converter._soffice_bin', return_value=None):
            self.assertEqual(convert_docx_to_pdf(self.docx, out), self.docx)

        def successful_run(*args, **kwargs):
            (out / 'source.pdf').write_bytes(b'%PDF')
            return SimpleNamespace(returncode=0, stdout='ok', stderr='')

        with patch('documents.pdf_converter._soffice_bin', return_value='/usr/bin/soffice'), \
             patch('documents.pdf_converter.subprocess.run', side_effect=successful_run):
            self.assertEqual(_convert_via_spawn(self.docx, out), out / 'source.pdf')
        (out / 'source.pdf').unlink()
        with patch('documents.pdf_converter._soffice_bin', return_value='/usr/bin/soffice'), \
             patch('documents.pdf_converter.subprocess.run', return_value=SimpleNamespace(returncode=1, stdout='', stderr='bad')):
            self.assertEqual(_convert_via_spawn(self.docx, out), self.docx)
        with patch('documents.pdf_converter._soffice_bin', return_value='/usr/bin/soffice'), \
             patch('documents.pdf_converter.subprocess.run', side_effect=__import__('subprocess').TimeoutExpired('x', 1)):
            self.assertEqual(_convert_via_spawn(self.docx, out), self.docx)

    @override_settings(DOCUMENT_PDF_ENABLED=True, DOCUMENT_PDF_BACKEND='uno')
    def test_uno_dispatch_falls_back_to_spawn(self):
        from documents.pdf_converter import convert_docx_to_pdf

        with patch('documents.pdf_converter._convert_via_uno', return_value=self.docx) as uno:
            self.assertEqual(convert_docx_to_pdf(self.docx), self.docx)
        uno.assert_called_once()


@override_settings(STORAGES=_TEST_STORAGES)
class DocumentViewCoverageTests(TestCase):
    """Cover document authorization, caching, and admin-managed content."""

    def setUp(self):
        from accounts.models import User
        from core.models import Request, Service
        from documents.models import DocumentBlock, ServiceTemplate

        self.media_root = tempfile.mkdtemp(prefix='plagenor-document-views-')
        self.settings_override = override_settings(
            MEDIA_ROOT=self.media_root, DOCUMENT_PDF_ENABLED=False,
        )
        self.settings_override.enable()
        self.admin = User.objects.create_user(
            username='doc-admin', password='pass', role='SUPER_ADMIN',
            is_staff=True, is_superuser=True,
        )
        self.owner = User.objects.create_user(
            username='doc-owner', password='pass', role='REQUESTER',
        )
        self.other = User.objects.create_user(
            username='doc-other', password='pass', role='REQUESTER',
        )
        analyst_user = User.objects.create_user(
            username='doc-member', password='pass', role='MEMBER',
        )
        self.service = Service.objects.create(
            code='DOC-VIEW', name='Document view service',
            channel_availability='BOTH',
        )
        self.req = Request.objects.create(
            display_id='DOC-VIEW-1', title='Document view', channel='IBTIKAR',
            status='ASSIGNED', service=self.service, requester=self.owner,
            assigned_to=analyst_user.member_profile,
        )
        upload = SimpleUploadedFile(
            'template.docx', b'placeholder docx',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )
        self.template = ServiceTemplate.objects.create(
            service=self.service, template_type='QUOTE', name='Initial template',
            file=upload, created_by=self.admin, is_active=True,
        )
        self.block = DocumentBlock.objects.create(
            template_type='QUOTE', position='BOTTOM', language='fr',
            title='Notice', body='Initial body', created_by=self.admin,
            updated_by=self.admin,
        )

    def tearDown(self):
        self.settings_override.disable()
        shutil.rmtree(self.media_root, ignore_errors=True)

    def test_download_permissions_guest_and_cache(self):
        from documents.views import (
            _block_signature, _cached_serve_doc, _service_fields_signature,
        )

        source = Path(self.media_root) / 'generated.docx'
        doc = _OpenDocx()
        doc.add_paragraph('generated content')
        doc.save(source)
        generator = MagicMock(return_value=str(source))

        self.client.force_login(self.other)
        denied = self.client.get(f'/documents/quote/{self.req.pk}/')
        self.assertEqual(denied.status_code, 403)
        self.client.force_login(self.owner)
        with patch('documents.views.generate_quote', generator):
            allowed = self.client.get(f'/documents/quote/{self.req.pk}/')
        self.assertEqual(allowed.status_code, 200)
        self.assertIn('attachment;', allowed['Content-Disposition'])

        # A cached second response must not regenerate the file.
        generator.reset_mock()
        response = _cached_serve_doc(self.req, 'QUOTE', generator, 'quote')
        self.assertEqual(response.status_code, 200)
        generator.assert_not_called()
        self.assertNotEqual(_block_signature(self.req, 'QUOTE'), '0')
        self.assertEqual(_service_fields_signature(SimpleNamespace(service_id=None)), '0')

        self.req.submitted_as_guest = True
        self.req.guest_token = __import__('uuid').uuid4()
        self.req.save(update_fields=['submitted_as_guest', 'guest_token'])
        guest_source = Path(self.media_root) / 'guest.docx'
        guest_doc = _OpenDocx()
        guest_doc.add_paragraph('guest content')
        guest_doc.save(guest_source)
        with patch('documents.views.generate_ibtikar_form', return_value=str(guest_source)):
            guest = self.client.get(
                f'/documents/guest/ibtikar-form/{self.req.guest_token}/')
        self.assertEqual(guest.status_code, 200)

    def test_template_management_crud_and_toggle(self):
        from documents.models import ServiceTemplate

        self.client.force_login(self.admin)
        for url in (
            '/documents/templates/',
            f'/documents/templates/{self.template.pk}/',
            f'/documents/templates/{self.template.pk}/edit/',
            f'/documents/templates/{self.template.pk}/delete/',
        ):
            self.assertEqual(self.client.get(url).status_code, 200)

        invalid = self.client.post('/documents/templates/create/', {'name': ''})
        self.assertEqual(invalid.status_code, 200)
        created = self.client.post('/documents/templates/create/', {
            'service': str(self.service.pk), 'template_type': 'IBTIKAR_FORM',
            'name': 'Uploaded template', 'description': 'Coverage',
            'file': SimpleUploadedFile('new.docx', b'docx'),
        })
        self.assertEqual(created.status_code, 302)
        uploaded = ServiceTemplate.objects.get(name='Uploaded template')

        self.client.post(f'/documents/templates/{uploaded.pk}/toggle/')
        uploaded.refresh_from_db()
        self.assertFalse(uploaded.is_active)
        self.client.post(f'/documents/templates/{uploaded.pk}/toggle/')
        uploaded.refresh_from_db()
        self.assertTrue(uploaded.is_active)

        invalid_edit = self.client.post(
            f'/documents/templates/{uploaded.pk}/edit/', {'name': ''})
        self.assertEqual(invalid_edit.status_code, 200)
        valid_edit = self.client.post(f'/documents/templates/{uploaded.pk}/edit/', {
            'name': 'Renamed template', 'description': 'Updated', 'is_active': 'on',
            'file': SimpleUploadedFile('replacement.docx', b'new docx'),
        })
        self.assertEqual(valid_edit.status_code, 302)
        deleted = self.client.post(f'/documents/templates/{uploaded.pk}/delete/')
        self.assertEqual(deleted.status_code, 302)
        self.assertFalse(ServiceTemplate.objects.filter(pk=uploaded.pk).exists())

    def test_document_block_crud_filters_and_validation(self):
        from documents.models import DocumentBlock

        self.client.force_login(self.admin)
        for query in ('', '?type=QUOTE&service=global&lang=fr',
                      f'?service={self.service.pk}'):
            self.assertEqual(self.client.get('/documents/blocks/' + query).status_code, 200)
        self.assertEqual(self.client.get('/documents/blocks/create/').status_code, 200)

        invalid_payloads = [
            {'template_type': 'BAD', 'position': 'BOTTOM', 'language': 'fr', 'body': 'x'},
            {'template_type': 'QUOTE', 'position': 'BAD', 'language': 'fr', 'body': 'x'},
            {'template_type': 'QUOTE', 'position': 'BOTTOM', 'language': 'xx', 'body': 'x'},
            {'template_type': 'QUOTE', 'position': 'BOTTOM', 'language': 'fr', 'body': ''},
            {'template_type': 'QUOTE', 'position': 'BOTTOM', 'language': 'fr',
             'body': 'x', 'services': ['not-a-uuid']},
        ]
        for payload in invalid_payloads:
            self.assertEqual(self.client.post('/documents/blocks/create/', payload).status_code, 200)

        created = self.client.post('/documents/blocks/create/', {
            'template_type': 'QUOTE', 'position': 'TOP', 'language': 'en',
            'title': 'English notice', 'body': 'Required text', 'priority': 'bad',
            'is_active': 'on', 'services': [str(self.service.pk)],
        })
        self.assertEqual(created.status_code, 302)
        block = DocumentBlock.objects.get(title='English notice')
        self.assertFalse(block.is_global)
        self.assertIn(self.service.code, block.scope_label())
        self.assertEqual(
            list(DocumentBlock.applicable_blocks('QUOTE', self.service, 'en')),
            [block],
        )

        self.assertEqual(self.client.get(f'/documents/blocks/{block.pk}/edit/').status_code, 200)
        edited = self.client.post(f'/documents/blocks/{block.pk}/edit/', {
            'template_type': 'QUOTE', 'position': 'BOTTOM', 'language': 'fr',
            'title': 'Updated notice', 'body': 'Updated body', 'priority': '4',
        })
        self.assertEqual(edited.status_code, 302)
        self.client.post(f'/documents/blocks/{block.pk}/toggle/')
        block.refresh_from_db()
        self.assertTrue(block.is_active)
        self.assertEqual(self.client.get(f'/documents/blocks/{block.pk}/delete/').status_code, 200)
        self.assertEqual(self.client.post(f'/documents/blocks/{block.pk}/delete/').status_code, 302)
        self.assertFalse(DocumentBlock.objects.filter(pk=block.pk).exists())
