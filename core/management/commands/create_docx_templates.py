"""Management command to create physical DOCX template files."""
import os
from pathlib import Path
from django.core.management.base import BaseCommand
from django.conf import settings


class Command(BaseCommand):
    help = 'Create DOCX template files in documents/docx_templates/'

    def handle(self, *args, **options):
        try:
            from docx import Document
            from docx.shared import Inches, Pt, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            self.stdout.write(self.style.WARNING(
                'python-docx not installed. Installing...'
            ))
            os.system('pip install python-docx')
            from docx import Document
            from docx.shared import Inches, Pt, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH

        output_dir = settings.BASE_DIR / 'documents' / 'docx_templates'
        output_dir.mkdir(parents=True, exist_ok=True)

        self._create_ibtikar_form(Document, Pt, WD_ALIGN_PARAGRAPH, output_dir)
        self._create_platform_note(Document, Pt, WD_ALIGN_PARAGRAPH, output_dir)
        self._create_reception_form(Document, Pt, WD_ALIGN_PARAGRAPH, output_dir)
        self._create_quote_template(Document, Pt, WD_ALIGN_PARAGRAPH, output_dir)

        self.stdout.write(self.style.SUCCESS(
            f'Created 4 DOCX templates in {output_dir}'
        ))

    def _create_ibtikar_form(self, Document, Pt, WD_ALIGN, output_dir):
        doc = Document()
        style = doc.styles['Normal']
        style.font.size = Pt(11)

        # Title
        title = doc.add_heading('PLAGENOR — Formulaire IBTIKAR', level=1)
        title.alignment = WD_ALIGN.CENTER
        subtitle = doc.add_paragraph('ESSBO — École Supérieure en Sciences Biologiques d\'Oran')
        subtitle.alignment = WD_ALIGN.CENTER

        doc.add_paragraph('')

        # Requester info table
        doc.add_heading('Informations du demandeur', level=2)
        table = doc.add_table(rows=7, cols=2)
        table.style = 'Table Grid'
        fields = [
            ('Nom et prénom', '{{FULL_NAME}}'),
            ('Établissement', '{{ETABLISSEMENT}}'),
            ('Laboratoire', '{{LABORATORY}}'),
            ('Titre du projet', '{{PROJECT_TITLE}}'),
            ('Directeur de recherche', '{{SUPERVISOR}}'),
            ('Téléphone', '{{PHONE}}'),
            ('Email', '{{EMAIL}}'),
        ]
        for i, (label, placeholder) in enumerate(fields):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = placeholder

        doc.add_paragraph('')

        # Service section
        doc.add_heading('Détails du service', level=2)
        table2 = doc.add_table(rows=3, cols=2)
        table2.style = 'Table Grid'
        table2.cell(0, 0).text = 'Service demandé'
        table2.cell(0, 1).text = '{{SERVICE_NAME}}'
        table2.cell(1, 0).text = 'Référence'
        table2.cell(1, 1).text = '{{DISPLAY_ID}}'
        table2.cell(2, 0).text = 'Date'
        table2.cell(2, 1).text = '{{DATE}}'

        doc.add_paragraph('')

        # Sample table placeholder
        doc.add_heading('Tableau des échantillons', level=2)
        doc.add_paragraph('[Tableau des échantillons à remplir]')

        doc.add_paragraph('')
        doc.add_paragraph('')

        # Signature
        sig = doc.add_paragraph('Signature du demandeur: ________________________')
        sig.alignment = WD_ALIGN.RIGHT
        date_sig = doc.add_paragraph('Date: ________________________')
        date_sig.alignment = WD_ALIGN.RIGHT

        doc.save(str(output_dir / 'ibtikar_form_template.docx'))

    def _create_platform_note(self, Document, Pt, WD_ALIGN, output_dir):
        doc = Document()
        style = doc.styles['Normal']
        style.font.size = Pt(11)

        title = doc.add_heading('Note de Plateforme — PLAGENOR', level=1)
        title.alignment = WD_ALIGN.CENTER

        doc.add_paragraph('')
        doc.add_paragraph('Référence: {{DISPLAY_ID}}')
        doc.add_paragraph('')

        doc.add_heading('Informations étudiant', level=2)
        table = doc.add_table(rows=2, cols=2)
        table.style = 'Table Grid'
        table.cell(0, 0).text = 'Nom et prénom'
        table.cell(0, 1).text = '{{FULL_NAME}}'
        table.cell(1, 0).text = 'Établissement'
        table.cell(1, 1).text = '{{ETABLISSEMENT}}'

        doc.add_paragraph('')

        doc.add_heading('Service', level=2)
        doc.add_paragraph('Service: {{SERVICE_NAME}}')

        doc.add_paragraph('')

        doc.add_heading('Déduction budgétaire', level=2)
        doc.add_paragraph('Montant: {{BUDGET_AMOUNT}} / 200 000 DA')

        doc.add_paragraph('')
        sig = doc.add_paragraph('Visa de la plateforme: ________________________')
        sig.alignment = WD_ALIGN.RIGHT

        doc.save(str(output_dir / 'platform_note_template.docx'))

    def _create_reception_form(self, Document, Pt, WD_ALIGN, output_dir):
        doc = Document()
        style = doc.styles['Normal']
        style.font.size = Pt(11)

        title = doc.add_heading('Fiche de Réception d\'Échantillons', level=1)
        title.alignment = WD_ALIGN.CENTER

        doc.add_paragraph('')

        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
        table.cell(0, 0).text = 'Date RDV'
        table.cell(0, 1).text = '{{APPOINTMENT_DATE}}'
        table.cell(1, 0).text = 'Code suivi'
        table.cell(1, 1).text = '{{TRACKING_CODE}}'
        table.cell(2, 0).text = 'Demandeur'
        table.cell(2, 1).text = '{{FULL_NAME}}'

        doc.add_paragraph('')

        doc.add_heading('Détails des échantillons', level=2)
        doc.add_paragraph('[Détails des échantillons reçus]')

        doc.add_paragraph('')

        doc.add_heading('État des échantillons', level=2)
        doc.add_paragraph('[ ] Conforme   [ ] Non conforme')
        doc.add_paragraph('Observations: ____________________________________________')

        doc.add_paragraph('')
        sig = doc.add_paragraph('Réceptionné par: ________________________')
        sig.alignment = WD_ALIGN.RIGHT

        doc.save(str(output_dir / 'reception_form_template.docx'))

    def _create_quote_template(self, Document, Pt, WD_ALIGN, output_dir):
        doc = Document()
        style = doc.styles['Normal']
        style.font.size = Pt(11)

        title = doc.add_heading('Devis — GENOCLAB', level=1)
        title.alignment = WD_ALIGN.CENTER

        doc.add_paragraph('')
        doc.add_paragraph('Client: {{CLIENT_NAME}}')
        doc.add_paragraph('')

        doc.add_heading('Lignes de facturation', level=2)
        table = doc.add_table(rows=2, cols=4)
        table.style = 'Table Grid'
        table.cell(0, 0).text = 'Description'
        table.cell(0, 1).text = 'Quantité'
        table.cell(0, 2).text = 'Prix unitaire'
        table.cell(0, 3).text = 'Total'
        table.cell(1, 0).text = '[Service]'
        table.cell(1, 1).text = '[Qté]'
        table.cell(1, 2).text = '[Prix]'
        table.cell(1, 3).text = '[Total]'

        doc.add_paragraph('')

        # Totals
        totals = doc.add_table(rows=3, cols=2)
        totals.style = 'Table Grid'
        totals.cell(0, 0).text = 'Sous-total HT'
        totals.cell(0, 1).text = '{{SUBTOTAL_HT}} DA'
        totals.cell(1, 0).text = 'TVA (19%)'
        totals.cell(1, 1).text = '{{TVA}} DA'
        totals.cell(2, 0).text = 'Total TTC'
        totals.cell(2, 1).text = '{{TOTAL_TTC}} DA'

        doc.add_paragraph('')
        sig = doc.add_paragraph('Signature: ________________________')
        sig.alignment = WD_ALIGN.RIGHT

        doc.save(str(output_dir / 'quote_template.docx'))
