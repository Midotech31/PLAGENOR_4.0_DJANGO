"""Rendu du rapport d'évaluation en DOCX (python-docx) et PDF (ReportLab).

Aucune ressource distante : polices intégrées aux moteurs, aucune image externe.
"""

from __future__ import annotations

from io import BytesIO

from app.core.config import DRAFT_BANNER, SIGNATURE
from app.reports.evaluation_report import Block, EvaluationReport

#: Couleurs de la direction artistique (bleu nuit, vert profond, ambre, rouge).
BLEU_NUIT = "123342"
VERT_PROFOND = "176B5B"
AMBRE = "B97812"
ROUGE = "B33A3A"
GRIS = "F5F7F6"

TONE_COLOR = {"neutre": BLEU_NUIT, "attention": AMBRE, "critique": ROUGE}

#: Étiquette compacte accolée à chaque paragraphe factuel.
KIND_LABEL = {
    "FAIT_EXTRAIT": "FAIT EXTRAIT",
    "CALCUL": "CALCUL",
    "ALERTE_SYSTEME": "ALERTE SYSTÈME",
    "COMMENTAIRE_EVALUATEUR": "COMMENTAIRE ÉVALUATEUR",
    "CONCLUSION_EVALUATEUR": "CONCLUSION ÉVALUATEUR",
    "A_VERIFIER": "À VÉRIFIER",
}

DRAFT_WARNING = (
    "BROUILLON NON VALIDÉ — la porte G7_VALIDATION_HUMAINE n'est pas satisfaite. "
    "Ce document ne peut pas être utilisé comme rapport officiel."
)


def _annotation(block: Block) -> str:
    """Suffixe de traçabilité : étiquette + page source."""
    paragraph = block.paragraph
    assert paragraph is not None
    label = KIND_LABEL.get(paragraph.kind, paragraph.kind)
    source = paragraph.source_label
    return f"[{label}{' — ' + source if source else ''}]"


# --------------------------------------------------------------------------
# DOCX
# --------------------------------------------------------------------------


def write_docx(model: EvaluationReport) -> bytes:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    document = Document()
    document.core_properties.title = f"Rapport d'évaluation — {model.reference}"
    document.core_properties.author = model.evaluator
    document.core_properties.comments = SIGNATURE

    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    def shade(cell, color: str) -> None:
        element = OxmlElement("w:shd")
        element.set(qn("w:fill"), color)
        cell._tc.get_or_add_tcPr().append(element)

    # -- bandeau ------------------------------------------------------------
    banner = document.add_paragraph()
    banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = banner.add_run(DRAFT_BANNER)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(AMBRE)

    if model.is_draft:
        warning = document.add_paragraph()
        warning.alignment = WD_ALIGN_PARAGRAPH.CENTER
        warning_run = warning.add_run(DRAFT_WARNING)
        warning_run.italic = True
        warning_run.font.size = Pt(9)
        warning_run.font.color.rgb = RGBColor.from_string(ROUGE)

    # -- titre --------------------------------------------------------------
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("RAPPORT D'ÉVALUATION")
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor.from_string(BLEU_NUIT)

    for text, size in ((model.subtitle, 11), (model.title, 12), (model.organizer, 10)):
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        run.font.size = Pt(size)
        if size == 12:
            run.bold = True

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(
        f"Référence {model.reference} — évaluateur {model.evaluator} — "
        f"généré le {model.generated_at.strftime('%d/%m/%Y à %H:%M UTC')} — "
        f"version {model.version}"
    )
    meta_run.font.size = Pt(8)

    # -- encadré d'avis proposé, en tête du rapport --------------------------
    if model.headline is not None:
        document.add_paragraph()
        headline = document.add_table(rows=1, cols=1)
        headline.style = "Table Grid"
        headline.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = headline.rows[0].cells[0]
        shade(cell, GRIS)
        cell.text = ""
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(model.headline.title + ". ")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor.from_string(
            TONE_COLOR.get(model.headline.tone, BLEU_NUIT)
        )
        paragraph.add_run(model.headline.body)
        document.add_paragraph()

    # -- corps --------------------------------------------------------------
    for section in model.sections:
        heading = document.add_heading(level=1)
        heading_run = heading.add_run(f"{section.number}. {section.title}")
        heading_run.font.color.rgb = RGBColor.from_string(BLEU_NUIT)

        for block in section.blocks:
            if block.kind == "box" and block.box is not None:
                table = document.add_table(rows=1, cols=1)
                table.style = "Table Grid"
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                cell = table.rows[0].cells[0]
                shade(cell, GRIS)
                cell.text = ""
                paragraph = cell.paragraphs[0]
                title_run = paragraph.add_run(block.box.title + ". ")
                title_run.bold = True
                title_run.font.color.rgb = RGBColor.from_string(
                    TONE_COLOR.get(block.box.tone, BLEU_NUIT)
                )
                paragraph.add_run(block.box.body)
                document.add_paragraph()

            elif block.kind == "subheading":
                sub = document.add_heading(level=2)
                sub_run = sub.add_run(block.text)
                sub_run.font.color.rgb = RGBColor.from_string(VERT_PROFOND)

            elif block.kind == "paragraph" and block.paragraph is not None:
                paragraph = document.add_paragraph()
                paragraph.add_run(block.paragraph.text)
                marker = paragraph.add_run("  " + _annotation(block))
                marker.font.size = Pt(7)
                marker.italic = True
                marker.font.color.rgb = RGBColor.from_string(VERT_PROFOND)

            elif block.kind == "list":
                for item in block.items:
                    document.add_paragraph(item, style="List Bullet")

            elif block.kind == "table" and block.table is not None:
                spec = block.table
                table = document.add_table(rows=1, cols=len(spec.headers))
                table.style = "Table Grid"
                for index, header in enumerate(spec.headers):
                    cell = table.rows[0].cells[index]
                    cell.text = ""
                    run = cell.paragraphs[0].add_run(header)
                    run.bold = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor.from_string(BLEU_NUIT)
                    shade(cell, GRIS)
                for row in spec.rows:
                    cells = table.add_row().cells
                    for index, value in enumerate(row[: len(spec.headers)]):
                        cells[index].text = ""
                        run = cells[index].paragraphs[0].add_run(str(value))
                        run.font.size = Pt(8.5)
                if spec.note:
                    note = document.add_paragraph()
                    note_run = note.add_run(spec.note)
                    note_run.italic = True
                    note_run.font.size = Pt(7.5)
                document.add_paragraph()

    # -- signature ----------------------------------------------------------
    document.add_paragraph()
    signature = document.add_paragraph()
    signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    signature_run = signature.add_run(SIGNATURE)
    signature_run.italic = True
    signature_run.font.color.rgb = RGBColor.from_string(BLEU_NUIT)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


# --------------------------------------------------------------------------
# PDF
# --------------------------------------------------------------------------


def write_pdf(model: EvaluationReport) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_RIGHT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        ListFlowable,
        ListItem,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    # En densité « compact », marges, corps et interlignes sont resserrés au
    # maximum de ce que la lisibilité permet. Le contenu, lui, est identique :
    # rien n'est jamais retiré pour gagner une page.
    compact = model.density == "compact"
    margin = 11 if compact else 16
    scale = 0.82 if compact else 1.0
    pad = 2 if compact else 3

    buffer = BytesIO()
    document = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        title=f"Rapport d'évaluation — {model.reference}",
        author=model.evaluator,
        subject=SIGNATURE,
        leftMargin=margin * mm,
        rightMargin=margin * mm,
        topMargin=(10 if compact else 16) * mm,
        bottomMargin=(13 if compact else 18) * mm,
    )

    base = getSampleStyleSheet()
    navy = colors.HexColor("#" + BLEU_NUIT)
    green = colors.HexColor("#" + VERT_PROFOND)

    st_banner = ParagraphStyle("banner", parent=base["Title"], fontSize=12, alignment=TA_CENTER,
                               textColor=colors.HexColor("#" + AMBRE), spaceAfter=2)
    st_warn = ParagraphStyle("warn", parent=base["Normal"], fontSize=8, alignment=TA_CENTER,
                             textColor=colors.HexColor("#" + ROUGE), spaceAfter=6)
    st_title = ParagraphStyle("title", parent=base["Title"], fontSize=18 * scale,
                              alignment=TA_CENTER, textColor=navy, spaceAfter=2)
    st_sub = ParagraphStyle("sub", parent=base["Normal"], fontSize=10.5, alignment=TA_CENTER,
                            textColor=navy, spaceAfter=1)
    st_meta = ParagraphStyle("meta", parent=base["Normal"], fontSize=7.5, alignment=TA_CENTER,
                             textColor=colors.grey, spaceAfter=10)
    st_h1 = ParagraphStyle("h1", parent=base["Heading1"], fontSize=12.5 * scale, textColor=navy,
                           spaceBefore=12 * scale, spaceAfter=5 * scale)
    st_h2 = ParagraphStyle("h2", parent=base["Heading2"], fontSize=10.5 * scale, textColor=green,
                           spaceBefore=8 * scale, spaceAfter=4 * scale)
    st_body = ParagraphStyle("body", parent=base["Normal"], fontSize=8.6 * scale,
                             leading=11.6 * scale, spaceAfter=3 * scale)
    st_tag = ParagraphStyle("tag", parent=base["Normal"], fontSize=6.6 * scale, textColor=green,
                            spaceAfter=5 * scale)
    st_cell = ParagraphStyle("cell", parent=base["Normal"], fontSize=7.6 * scale,
                             leading=9.3 * scale)
    st_head = ParagraphStyle("head", parent=base["Normal"], fontSize=7.6 * scale,
                             leading=9.3 * scale, textColor=colors.white,
                             fontName="Helvetica-Bold")
    st_note = ParagraphStyle("note", parent=base["Normal"], fontSize=6.8 * scale,
                             leading=8.6 * scale, textColor=colors.grey, spaceAfter=7 * scale)
    st_sign = ParagraphStyle("sign", parent=base["Normal"], alignment=TA_RIGHT,
                             fontName="Helvetica-Oblique", textColor=navy)

    def esc(value: str) -> str:
        return (
            str(value).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        )

    story: list = [Paragraph(esc(DRAFT_BANNER), st_banner)]
    if model.is_draft:
        story.append(Paragraph(esc(DRAFT_WARNING), st_warn))
    story += [
        Paragraph(esc(model.heading), st_title),
        Paragraph(esc(model.subtitle), st_sub),
        Paragraph(f"<b>{esc(model.title)}</b>", st_sub),
        Paragraph(esc(model.organizer), st_sub),
        Paragraph(
            esc(
                f"Référence {model.reference} — évaluateur {model.evaluator} — "
                f"généré le {model.generated_at.strftime('%d/%m/%Y à %H:%M UTC')} — "
                f"version {model.version}"
            ),
            st_meta,
        ),
    ]

    usable = A4[0] - 2 * margin * mm

    if model.headline is not None:
        tone_hex = TONE_COLOR.get(model.headline.tone, BLEU_NUIT)
        headline = Table(
            [[Paragraph(
                f'<font color="#{tone_hex}"><b>{esc(model.headline.title)}.</b></font> '
                f"{esc(model.headline.body)}",
                st_body,
            )]],
            colWidths=[usable],
        )
        headline.setStyle(
            TableStyle([
                ("BOX", (0, 0), (-1, -1), 1.4, colors.HexColor("#" + tone_hex)),
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#" + GRIS)),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 7),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
            ])
        )
        story += [headline, Spacer(1, 10 * scale)]

    for section in model.sections:
        story.append(Paragraph(esc(f"{section.number}. {section.title}"), st_h1))

        for block in section.blocks:
            if block.kind == "box" and block.box is not None:
                tone = colors.HexColor("#" + TONE_COLOR.get(block.box.tone, BLEU_NUIT))
                content = Paragraph(
                    f'<font color="#{TONE_COLOR.get(block.box.tone, BLEU_NUIT)}"><b>'
                    f"{esc(block.box.title)}.</b></font> {esc(block.box.body)}",
                    st_body,
                )
                boxed = Table([[content]], colWidths=[usable])
                boxed.setStyle(
                    TableStyle(
                        [
                            ("BOX", (0, 0), (-1, -1), 1.1, tone),
                            ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#" + GRIS)),
                            ("LEFTPADDING", (0, 0), (-1, -1), 7),
                            ("RIGHTPADDING", (0, 0), (-1, -1), 7),
                            ("TOPPADDING", (0, 0), (-1, -1), 6),
                            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                        ]
                    )
                )
                story += [boxed, Spacer(1, 7 * scale)]

            elif block.kind == "subheading":
                story.append(Paragraph(esc(block.text), st_h2))

            elif block.kind == "paragraph" and block.paragraph is not None:
                story.append(Paragraph(esc(block.paragraph.text), st_body))
                story.append(Paragraph(esc(_annotation(block)), st_tag))

            elif block.kind == "list":
                story.append(
                    ListFlowable(
                        [ListItem(Paragraph(esc(item), st_body), leftIndent=10) for item in block.items],
                        bulletType="bullet",
                        start="•",
                        leftIndent=12,
                    )
                )
                story.append(Spacer(1, 5 * scale))

            elif block.kind == "table" and block.table is not None:
                spec = block.table
                count = len(spec.headers)
                # Première colonne plus large : elle porte le libellé.
                if spec.widths and len(spec.widths) == count:
                    total = sum(spec.widths) or 1.0
                    widths = [usable * (part / total) for part in spec.widths]
                elif count == 1:
                    widths = [usable]
                else:
                    first = usable * (0.30 if count <= 3 else 0.24)
                    widths = [first] + [(usable - first) / (count - 1)] * (count - 1)

                data = [[Paragraph(esc(header), st_head) for header in spec.headers]]
                for row in spec.rows:
                    data.append(
                        [Paragraph(esc(value), st_cell) for value in row[:count]]
                        + [Paragraph("", st_cell)] * max(0, count - len(row))
                    )

                table = Table(data, colWidths=widths, repeatRows=1)
                table.setStyle(
                    TableStyle(
                        [
                            ("BACKGROUND", (0, 0), (-1, 0), navy),
                            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#D9E0E2")),
                            ("VALIGN", (0, 0), (-1, -1), "TOP"),
                            ("ROWBACKGROUNDS", (0, 1), (-1, -1),
                             [colors.white, colors.HexColor("#FAFCFB")]),
                            ("LEFTPADDING", (0, 0), (-1, -1), pad + 1),
                            ("RIGHTPADDING", (0, 0), (-1, -1), pad + 1),
                            ("TOPPADDING", (0, 0), (-1, -1), pad),
                            ("BOTTOMPADDING", (0, 0), (-1, -1), pad),
                        ]
                    )
                )
                story.append(table)
                if spec.note:
                    story.append(Paragraph(esc(spec.note), st_note))
                story.append(Spacer(1, 6 * scale))

    story += [Spacer(1, 12 * scale), Paragraph(esc(SIGNATURE), st_sign)]

    def decorate(canvas, doc) -> None:
        canvas.saveState()
        canvas.setFont("Helvetica", 6.5)
        canvas.setFillColor(colors.grey)
        canvas.drawString(16 * mm, 10 * mm, DRAFT_BANNER)
        canvas.drawCentredString(A4[0] / 2, 10 * mm, f"Page {doc.page}")
        canvas.drawRightString(A4[0] - 16 * mm, 10 * mm, SIGNATURE)
        if model.is_draft:
            canvas.saveState()
            canvas.setFont("Helvetica-Bold", 60)
            canvas.setFillColorRGB(0.87, 0.87, 0.87)
            canvas.translate(A4[0] / 2, A4[1] / 2)
            canvas.rotate(45)
            canvas.drawCentredString(0, 0, "BROUILLON")
            canvas.restoreState()
        canvas.restoreState()

    document.build(story, onFirstPage=decorate, onLaterPages=decorate)
    return buffer.getvalue()


__all__ = ["write_docx", "write_pdf"]
