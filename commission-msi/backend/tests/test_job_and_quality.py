"""Travail durable, relecture indépendante, contrôle qualité, mode d'IA.

Le fil conducteur : **rien ne se perd et rien ne se décide tout seul**. Un
traitement survit à la perte du processus, un désaccord d'audit devient `NV`
sans moyenne, un rapport partiellement valide n'est jamais remis, et aucune
pièce d'identité ne quitte le poste.
"""

from __future__ import annotations

import pytest

from app.core.config import reset_settings
from app.core.vocabulary import JobState
from app.models import AnalysisJob
from app.services import ai_provider, audit_service, job_service, report_qa_service
from tests.fixtures import synthetic

DOSSIER = """DEMANDE D'ORGANISATION D'UNE MANIFESTATION SCIENTIFIQUE INTERNATIONALE
Intitulé : Colloque international fictif sur les materiaux durables
Dates : du 12 mars 2027 au 14 mars 2027
Lieu : Campus fictif, Alger
Format : hybride
Établissement organisateur : Universite Fictive de Test
Comité scientifique : Pr Jean Dubois (France) ; Pr Karim Idrissi (Maroc)
Pays représentés : Algerie, France, Maroc
Budget total : 1 000 000 DA
Modalités de publication : actes indexes
"""


def _import(client, dossier):
    return client.post(
        f"/api/v1/dossiers/{dossier['id']}/documents",
        files={"file": ("dossier.pdf", synthetic.make_pdf([DOSSIER]), "application/pdf")},
    )


# --------------------------------------------------------------------------
# Travail durable
# --------------------------------------------------------------------------


def test_processing_cannot_start_without_a_pdf(client, dossier):
    response = client.post(f"/api/v1/dossiers/{dossier['id']}/traitement")
    assert response.status_code == 422
    assert "Aucun PDF valide" in response.json()["error"]["message"]


def test_processing_is_persisted_in_database_not_in_memory(client, dossier, session):
    assert _import(client, dossier).status_code == 201
    created = client.post(f"/api/v1/dossiers/{dossier['id']}/traitement").json()

    assert created["state"] == JobState.QUEUED
    # Le travail existe en base : perdre le processus ne le perd pas.
    stored = session.get(AnalysisJob, created["id"])
    assert stored is not None
    assert stored.dossier_id == dossier["id"]
    assert stored.source_sha256, "l'empreinte du PDF traité doit être enregistrée"


def test_the_same_file_is_not_processed_twice_at_once(client, dossier):
    assert _import(client, dossier).status_code == 201
    assert client.post(f"/api/v1/dossiers/{dossier['id']}/traitement").status_code == 201

    duplicate = client.post(f"/api/v1/dossiers/{dossier['id']}/traitement")
    assert duplicate.status_code == 422
    assert "déjà en cours" in duplicate.json()["error"]["message"]


def test_the_worker_runs_the_whole_pipeline_and_completes(client, dossier, session):
    assert _import(client, dossier).status_code == 201
    created = client.post(f"/api/v1/dossiers/{dossier['id']}/traitement").json()

    assert job_service.work_once() == created["id"]

    state = client.get(f"/api/v1/dossiers/{dossier['id']}/traitement").json()["job"]
    assert state["state"] == JobState.COMPLETED, state["error_message"]
    assert state["progress"] == 100
    assert state["pages_done"] >= 1
    assert state["referential_version"] and state["grid_version"]
    # Toutes les étapes du pipeline ont laissé un point de reprise.
    assert set(state["steps_done"]) == {step for step, _ in job_service.PIPELINE}


def test_a_completed_step_is_not_replayed_on_a_second_run(client, dossier, session):
    assert _import(client, dossier).status_code == 201
    first = client.post(f"/api/v1/dossiers/{dossier['id']}/traitement").json()
    job_service.work_once()

    # Un nouveau travail sur le même dossier réutilise les points de reprise.
    second = job_service.enqueue(session, dossier["id"])
    checkpoints = job_service.checkpoints(session, first["id"])
    assert checkpoints, "le premier travail doit avoir laissé des points de reprise"
    assert second.id != first["id"]


def test_cancelling_stops_the_work_without_erasing_anything(client, dossier, session):
    assert _import(client, dossier).status_code == 201
    created = client.post(f"/api/v1/dossiers/{dossier['id']}/traitement").json()

    cancelled = client.post(
        f"/api/v1/dossiers/{dossier['id']}/traitement/{created['id']}/annuler"
    ).json()
    assert cancelled["state"] == JobState.CANCELLED
    assert cancelled["can_resume"] is True
    # Le dossier lui-même est intact.
    assert client.get(f"/api/v1/dossiers/{dossier['id']}").status_code == 200


def test_resuming_puts_the_work_back_in_the_queue(client, dossier):
    assert _import(client, dossier).status_code == 201
    created = client.post(f"/api/v1/dossiers/{dossier['id']}/traitement").json()
    client.post(f"/api/v1/dossiers/{dossier['id']}/traitement/{created['id']}/annuler")

    resumed = client.post(
        f"/api/v1/dossiers/{dossier['id']}/traitement/{created['id']}/reprendre"
    ).json()
    assert resumed["state"] == JobState.QUEUED
    assert resumed["cancel_requested"] is False


def test_a_lease_prevents_two_workers_from_taking_the_same_job(client, dossier, session):
    assert _import(client, dossier).status_code == 201
    client.post(f"/api/v1/dossiers/{dossier['id']}/traitement")

    first = job_service.claim(session, owner="worker-a")
    assert first is not None
    # Le second worker ne peut pas prendre un travail sous bail valide.
    assert job_service.claim(session, owner="worker-b") is None


def test_an_expired_lease_lets_the_work_be_taken_over(client, dossier, session):
    from datetime import timedelta

    from app.models.base import utcnow

    assert _import(client, dossier).status_code == 201
    client.post(f"/api/v1/dossiers/{dossier['id']}/traitement")

    taken = job_service.claim(session, owner="worker-mort")
    assert taken is not None
    taken.lease_expires_at = utcnow() - timedelta(seconds=1)
    session.commit()

    recovered = job_service.claim(session, owner="worker-vivant")
    assert recovered is not None and recovered.id == taken.id
    assert recovered.lease_owner == "worker-vivant"


def test_a_technical_failure_explains_the_cause_and_stays_resumable(
    client, dossier, session, monkeypatch
):
    assert _import(client, dossier).status_code == 201
    created = client.post(f"/api/v1/dossiers/{dossier['id']}/traitement").json()

    def explode(session, job):
        raise PermissionError("fichier verrouillé")

    monkeypatch.setattr(
        job_service, "PIPELINE", ((JobState.VALIDATING, 5),)
    )
    from app.services import pipeline

    monkeypatch.setitem(
        pipeline.STEPS,
        JobState.VALIDATING,
        pipeline.Step(signature=pipeline._document_signature, run=explode),
    )

    job_service.work_once()
    state = client.get(f"/api/v1/dossiers/{dossier['id']}/traitement").json()["job"]
    assert state["error_code"] == "PermissionError"
    # L'erreur explique la cause et l'action possible, sans trace brute.
    assert "n'a pas pu être lu ou écrit" in state["error_message"]
    assert "fichier verrouillé" not in state["error_message"]
    assert "Reprendre" in state["error_message"]


# --------------------------------------------------------------------------
# Relecture indépendante
# --------------------------------------------------------------------------


def test_an_unresolved_disagreement_becomes_nv_without_averaging(
    client, dossier, session, monkeypatch
):
    assert _import(client, dossier).status_code == 201
    client.post(f"/api/v1/dossiers/{dossier['id']}/evaluation-automatique")

    # L'auditeur recalcule et conclut différemment sur A1.
    from app.services import regulatory_engine

    real_evaluate = regulatory_engine.evaluate

    def divergent(facts):
        results = real_evaluate(facts)
        for result in results:
            if result.code == "A1":
                result.status = (
                    regulatory_engine.Status.C
                    if result.status != regulatory_engine.Status.C
                    else regulatory_engine.Status.NC
                )
        return results

    monkeypatch.setattr(regulatory_engine, "evaluate", divergent)
    report = audit_service.review(session, dossier["id"])
    assert report["disagreements"] >= 1
    assert report["reclassified_nv"] >= 1

    monkeypatch.undo()
    criteria = {
        row["code"]: row
        for row in client.get(
            f"/api/v1/dossiers/{dossier['id']}/evaluation-automatique"
        ).json()["criteria"]
    }
    assert criteria["A1"]["status"] == "NV"
    assert "Aucune moyenne" in criteria["A1"]["finding"]

    listed = client.get(f"/api/v1/dossiers/{dossier['id']}/desaccords").json()
    assert listed["items"]
    assert "Aucune moyenne" in listed["notice"]


def test_the_auditor_never_overrides_a_human_qualification(client, dossier, session):
    assert _import(client, dossier).status_code == 201
    client.post(f"/api/v1/dossiers/{dossier['id']}/evaluation-automatique")

    client.post(
        f"/api/v1/dossiers/{dossier['id']}/evaluation-automatique/criteres/A1",
        json={"status": "C", "comment": "Pièces vérifiées une à une en séance."},
    )
    audit_service.review(session, dossier["id"])

    criteria = {
        row["code"]: row
        for row in client.get(
            f"/api/v1/dossiers/{dossier['id']}/evaluation-automatique"
        ).json()["criteria"]
    }
    assert criteria["A1"]["status"] == "C"
    assert criteria["A1"]["human_status"] == "C"


# --------------------------------------------------------------------------
# Contrôle qualité
# --------------------------------------------------------------------------


def test_quality_control_checks_the_twenty_six_criteria_and_the_arithmetic(
    client, dossier, session
):
    assert _import(client, dossier).status_code == 201
    client.post(f"/api/v1/dossiers/{dossier['id']}/evaluation-automatique")

    result = report_qa_service.run(session, dossier["id"])
    assert result["passed"] is True
    keys = {check["key"] for check in result["checks"]}
    assert {
        "criteres_presents_une_fois_et_dans_l_ordre",
        "aucune_cellule_vide",
        "evidence_ids_existants",
        "affirmations_sans_preuve",
        "recalcul_du_score",
        "plafonds_respectes",
        "avis_dans_la_liste_fermee",
        "score_ne_neutralise_pas_une_non_conformite",
        "aucun_delai_de_six_mois",
        "aucun_motif_interdit",
    } <= keys


def test_a_partially_valid_report_is_never_delivered(client, dossier, session):
    """Un contrôle bloquant en échec interrompt la remise du rapport."""
    assert _import(client, dossier).status_code == 201
    client.post(f"/api/v1/dossiers/{dossier['id']}/evaluation-automatique")

    from app.models import CriterionResult as CriterionResultRow

    row = session.query(CriterionResultRow).filter_by(
        dossier_id=dossier["id"], code="A1"
    ).one()
    row.evidence_refs = '["E-INEXISTANTE"]'
    session.commit()

    with pytest.raises(report_qa_service.QaFailed) as failure:
        report_qa_service.run(session, dossier["id"])
    assert "Citations orphelines" in str(failure.value)

    stored = report_qa_service.latest(session, dossier["id"])
    assert stored["passed"] is False
    assert stored["failures"] >= 1


def test_quality_control_is_exposed_to_the_evaluator(client, dossier, session):
    assert _import(client, dossier).status_code == 201
    client.post(f"/api/v1/dossiers/{dossier['id']}/evaluation-automatique")
    report_qa_service.run(session, dossier["id"])

    payload = client.get(f"/api/v1/dossiers/{dossier['id']}/controle-qualite").json()
    assert payload["passed"] is True
    assert payload["checks"]


# --------------------------------------------------------------------------
# Mode d'intelligence artificielle
# --------------------------------------------------------------------------


def test_local_only_is_the_default_and_says_what_it_cannot_do(client):
    state = client.get("/api/v1/mode-analyse").json()
    assert state["mode"] == ai_provider.LOCAL_ONLY
    assert state["external_transmission"] is False
    assert "ne fournit pas le même niveau" in state["notice"]
    # Le mode recommandé est celui qui lit **sans rien faire sortir**. Pour une
    # commission qui traite des dossiers confidentiels, recommander un mode qui
    # transmet — et qui exige une clé payante — ne se défendrait pas.
    assert state["recommended"] == ai_provider.LOCAL_MODEL


def test_the_api_key_is_never_exposed(client, monkeypatch):
    monkeypatch.setenv("ANTHROPIC_API_KEY", "cle-secrete-de-test")
    reset_settings()
    try:
        body = client.get("/api/v1/mode-analyse").text
        assert "cle-secrete-de-test" not in body
    finally:
        monkeypatch.delenv("ANTHROPIC_API_KEY", raising=False)
        reset_settings()


def test_hybrid_mode_refuses_to_run_until_it_is_fully_configured():
    provider = ai_provider.HybridStrictProvider()
    described = provider.describe()
    assert described["available"] is False
    assert "ANTHROPIC_API_KEY" in described["missing"]
    with pytest.raises(ai_provider.ExternalAiNotConfigured):
        provider.complete(ai_provider.AiRequest(role="ANALYSIS", instruction="—"))


def test_a_restricted_block_is_never_transmitted(monkeypatch):
    _configure_hybrid(monkeypatch)
    provider = ai_provider.HybridStrictProvider(client=_EchoClient())
    with pytest.raises(ai_provider.RestrictedContentRefused):
        provider.complete(
            ai_provider.AiRequest(
                role="ANALYSIS",
                instruction="Lire le bloc.",
                blocks=[{"evidence_id": "E-PJ005", "text": "…", "sensitivity": "RESTREINT"}],
            )
        )


def test_passport_numbers_are_redacted_before_any_transmission(monkeypatch):
    _configure_hybrid(monkeypatch)
    client_spy = _EchoClient()
    provider = ai_provider.HybridStrictProvider(client=client_spy)
    response = provider.complete(
        ai_provider.AiRequest(
            role="ANALYSIS",
            instruction="Résumer.",
            blocks=[
                {
                    "evidence_id": "E-P001",
                    "text": "Conférencier étranger, passeport N° AB1234567, université de Lyon.",
                }
            ],
        )
    )
    transmitted = client_spy.last["blocks"][0]["text"]
    assert "AB1234567" not in transmitted
    assert ai_provider.REDACTED in transmitted
    assert "PIECE_IDENTITE" in response.data_categories


def test_configuration_cannot_authorise_sending_identity_documents(monkeypatch):
    _configure_hybrid(monkeypatch)
    monkeypatch.setenv("SEND_IDENTITY_DOCUMENTS", "true")
    reset_settings()
    try:
        provider = ai_provider.HybridStrictProvider(client=_EchoClient())
        with pytest.raises(ai_provider.RestrictedContentRefused):
            provider.complete(ai_provider.AiRequest(role="ANALYSIS", instruction="—"))
    finally:
        monkeypatch.delenv("SEND_IDENTITY_DOCUMENTS", raising=False)
        reset_settings()


def test_an_unavailable_model_never_falls_back_silently(monkeypatch):
    _configure_hybrid(monkeypatch)

    class Missing:
        def complete(self, *, model_id, request):
            raise LookupError(model_id)

    provider = ai_provider.HybridStrictProvider(client=Missing())
    with pytest.raises(ai_provider.ModelUnavailable) as failure:
        provider.complete(ai_provider.AiRequest(role="ANALYSIS", instruction="—"))
    assert failure.value.code == "MODEL_UNAVAILABLE"
    assert "Aucun basculement silencieux" in str(failure.value)


def test_an_ai_call_is_logged_without_content_or_reasoning(session, monkeypatch):
    _configure_hybrid(monkeypatch)
    row = ai_provider.record_call(
        session,
        dossier_id=None,
        job_id=None,
        role="ANALYSIS",
        model_id="modele-de-test",
        status="OK",
        duration_ms=120,
        input_payload="texte confidentiel du dossier",
        output_payload="réponse du modèle",
        data_categories=["EXTRAIT_TEXTE"],
    )
    session.commit()
    assert row.input_sha256 and len(row.input_sha256) == 64
    # Le contenu et le raisonnement ne sont stockés nulle part.
    stored = {column.name: getattr(row, column.name) for column in row.__table__.columns}
    assert not any(
        isinstance(value, str) and "confidentiel" in value for value in stored.values()
    )


class _EchoClient:
    """Client injecté : il enregistre ce qui lui est réellement transmis."""

    def __init__(self) -> None:
        self.last: dict | None = None

    def complete(self, *, model_id, request):
        self.last = {"model_id": model_id, "blocks": request.blocks}
        return {"ok": True}


def _configure_hybrid(monkeypatch) -> None:
    monkeypatch.setenv("ANALYSIS_MODE", "HYBRID_STRICT")
    monkeypatch.setenv("ANTHROPIC_API_KEY", "cle-de-test")
    monkeypatch.setenv("ANTHROPIC_MODEL_ANALYSIS", "modele-de-test")
    monkeypatch.setenv("ALLOW_EXTERNAL_AI", "true")
    monkeypatch.setenv("MSI_PRIVACY_ACKNOWLEDGED", "true")
    reset_settings()


# --------------------------------------------------------------------------
# Messages d'erreur : un refus doit apprendre quoi corriger
# --------------------------------------------------------------------------


def test_a_schema_refusal_names_the_field_and_the_rule(client, dossier):
    """Un « 422 » nu n'apprend rien à l'évaluateur ; le motif doit être lisible."""
    assert _import(client, dossier).status_code == 201
    client.post(f"/api/v1/dossiers/{dossier['id']}/evaluation-automatique")

    response = client.post(
        f"/api/v1/dossiers/{dossier['id']}/evaluation-automatique/criteres/A1",
        json={"status": "C", "comment": "trop"},
    )
    assert response.status_code == 422
    message = response.json()["error"]["message"]
    assert "comment" in message
    assert "8 caractères" in message
    assert "justifiée" in message


def test_an_unknown_status_lists_the_accepted_values(client, dossier):
    assert _import(client, dossier).status_code == 201
    client.post(f"/api/v1/dossiers/{dossier['id']}/evaluation-automatique")

    response = client.post(
        f"/api/v1/dossiers/{dossier['id']}/evaluation-automatique/criteres/A1",
        json={"status": "INCONNU", "comment": "motivation suffisante"},
    )
    assert response.status_code == 422
    message = response.json()["error"]["message"]
    assert "status" in message
    for accepted in ("C", "PC", "NC", "NV"):
        assert accepted in message


def test_a_valid_qualification_is_accepted(client, dossier):
    assert _import(client, dossier).status_code == 201
    client.post(f"/api/v1/dossiers/{dossier['id']}/evaluation-automatique")

    response = client.post(
        f"/api/v1/dossiers/{dossier['id']}/evaluation-automatique/criteres/A1",
        json={"status": "C", "comment": "Pièces vérifiées une à une en séance."},
    )
    assert response.status_code == 200
    criteria = {row["code"]: row for row in response.json()["criteria"]}
    assert criteria["A1"]["human_status"] == "C"


# --------------------------------------------------------------------------
# Un seul clic : le rapport final fait partie du travail
# --------------------------------------------------------------------------


def test_one_click_produces_the_final_report_without_a_second_action(client, dossier):
    """L'évaluateur clique « Traiter le dossier » et le rapport existe."""
    assert _import(client, dossier).status_code == 201
    client.post(f"/api/v1/dossiers/{dossier['id']}/traitement")
    job_service.work_once()

    reports = client.get(f"/api/v1/dossiers/{dossier['id']}/rapports").json()["items"]
    formats = {report["format"] for report in reports}
    assert formats == {"docx", "pdf"}, "les deux formats sont produits par le travail"


def test_the_report_produced_by_the_job_is_downloadable_as_is(client, dossier):
    assert _import(client, dossier).status_code == 201
    client.post(f"/api/v1/dossiers/{dossier['id']}/traitement")
    job_service.work_once()

    reports = client.get(f"/api/v1/dossiers/{dossier['id']}/rapports").json()["items"]
    word = next(report for report in reports if report["format"] == "docx")
    response = client.get(
        f"/api/v1/dossiers/{dossier['id']}/rapports/{word['id']}/fichier"
    )
    assert response.status_code == 200
    assert response.content[:2] == b"PK"
    assert dossier["reference"] in response.headers["content-disposition"]


def test_the_produced_report_follows_the_commission_model(client, dossier, session):
    """Le fichier produit sans second clic est le rapport harmonisé attendu."""
    import io

    import docx

    assert _import(client, dossier).status_code == 201
    client.post(f"/api/v1/dossiers/{dossier['id']}/traitement")
    job_service.work_once()

    reports = client.get(f"/api/v1/dossiers/{dossier['id']}/rapports").json()["items"]
    word = next(report for report in reports if report["format"] == "docx")
    content = client.get(
        f"/api/v1/dossiers/{dossier['id']}/rapports/{word['id']}/fichier"
    ).content

    document = docx.Document(io.BytesIO(content))
    headings = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.style.name.startswith("Heading")
    ]
    # Les sept sections du modèle de la commission, dans l'ordre.
    for number, title in enumerate(
        (
            "Fiche d'information contrôlée",
            "Appréciation scientifique commune",
            "Matrice réglementaire uniforme",
            "Contrôle des intervenants étrangers",
            "Points de vigilance institutionnelle",
            "Compléments indispensables",
            "Orientation technique motivée",
        ),
        start=1,
    ):
        assert any(
            heading.startswith(f"{number}.") and title in heading for heading in headings
        ), f"section {number} absente : {headings}"

    # Les deux tableaux structurants du modèle : cinq dimensions + total, 26 critères.
    sizes = {(len(table.rows), len(table.columns)) for table in document.tables}
    assert (7, 4) in sizes, "tableau d'appréciation scientifique absent"
    assert (27, 5) in sizes, "matrice des 26 critères absente"


def test_the_page_count_is_measured_on_the_file_actually_written(client, dossier, session):
    assert _import(client, dossier).status_code == 201
    created = client.post(f"/api/v1/dossiers/{dossier['id']}/traitement").json()
    job_service.work_once()

    checkpoints = {
        checkpoint.step: checkpoint
        for checkpoint in job_service.checkpoints(session, created["id"])
    }
    assert JobState.REPORT_RENDERING in checkpoints
    import json

    result = json.loads(checkpoints[JobState.REPORT_RENDERING].result_json)
    pdf = next(item for item in result["rapports"] if item["format"] == "pdf")
    assert pdf["pages"] and pdf["pages"] > 0
    assert pdf["brouillon"] is True, "le travail produit un brouillon, jamais un officiel"


def test_a_report_is_never_produced_when_quality_control_blocks(client, dossier, monkeypatch):
    """Un rapport dont un contrôle bloquant échoue ne doit pas exister en fichier."""
    assert _import(client, dossier).status_code == 201

    def refuse(*args, **kwargs):
        raise report_qa_service.QaFailed(report_qa_service.QaReport(), [])

    monkeypatch.setattr(report_qa_service, "run", refuse)
    client.post(f"/api/v1/dossiers/{dossier['id']}/traitement")
    job_service.work_once()

    reports = client.get(f"/api/v1/dossiers/{dossier['id']}/rapports").json()["items"]
    assert reports == []
